"""Skeleton articles — the whole insight, rendered four ways, assembled into one .docx per persona.

For each persona we pick ONE anchor model and assemble: its charts, the persona's infographic (its most
apt layout family), a decorative Van Gogh header, and a 150–200-word grounded GIST. This is the
end-to-end preview the owner judges before full article wiring (hard-rule #1).

The gist is produced under the number-firewall (assessment §04 / CLAUDE.md #2): the LLM writes prose
but authors NO number — it cites executed values only by `{token}`, the renderer substitutes the
verified `NumberObject.rendered()`, and any figure the model typed that does not match an executed
value is caught (gate's leak regex) and re-narrated; on exhaustion we fall back to the persona's
authored, deterministically-filled `summary_template`. So every figure in the paragraph is provenance-
traced even though a .docx has no DOM gate.
"""
from __future__ import annotations

import base64
import re
from pathlib import Path

from pydantic import BaseModel, Field

from .illustration import vangogh
from .infographic.agentic import _tokenize
from .infographic.families import (cross_section_ladder, decision_brief, decomposition_hero,
                                   regime_dashboard)
from .infographic.from_persona import (chart_png, chart_png_family, clean_meaning, decisive, humanise,
                                      persona_material)
from .infographic.gate import _LEAK
from .studio.llm import get_llm

# The anchor model + layout family + (via vangogh.PERSONA_PAINTING) painting per persona. All four
# families appear across the eight; each family is used only where it is eligible.
MODEL_PICK: dict[str, str] = {
    "central_bank_policymaker": "reaction_function",
    "macro_rates_trader": "term_premium_decomposition",
    "volatility_trader": "garch_volatility",
    "credit_investor": "credit_quality_ladder",
    "equity_multiasset_pm": "financial_conditions",
    "commodity_analyst": "energy_complex",
    "corporate_treasurer": "funding_cost",
    "economist_forecaster": "beveridge_curve",
}
FAMILY = {
    "central_bank_policymaker": regime_dashboard,
    "macro_rates_trader": decomposition_hero,
    "volatility_trader": cross_section_ladder,
    "credit_investor": cross_section_ladder,
    "equity_multiasset_pm": regime_dashboard,
    "commodity_analyst": decision_brief,
    "corporate_treasurer": decomposition_hero,
    "economist_forecaster": regime_dashboard,
}
_TOKLEFT = re.compile(r"\{n\d+\}")


class Gist(BaseModel):
    body: str = Field(description="One flowing paragraph, 150-200 words, the decisive article read. "
                                  "Cite figures ONLY as the {n} tokens provided — never type a digit.")


# ── the 150–200-word gist, under the firewall ────────────────────────────────────────────────────
def _citable_for_model(mat: dict, model_id: str, limit: int = 8) -> tuple[dict, str]:
    """Token→NumberObject menu, the ANCHOR model's dimensioned numbers first, then the rest."""
    numbers = mat["numbers"]
    order = [k for k in mat["salient"] if k.startswith(f"{model_id}.")]
    order += [k for k in mat["salient"] if k not in order]
    keys = [k for k in order if any(c in "%$×σ°" or c.isalpha() for c in numbers[k].rendered())][:limit]
    toks, lines = {}, []
    for i, k in enumerate(keys):
        tok = f"n{i}"
        toks[tok] = numbers[k].model_copy(update={"name": tok})
        label = clean_meaning(mat["meanings"].get(k, ""), humanise(k.split(".", 1)[1]))
        lines.append(f"  {{{tok}}}  {label} = {numbers[k].rendered()}")
    return toks, "\n".join(lines)


def _narrate_gist(mat: dict, model_name: str, menu: str, feedback: str = "") -> str:
    p = mat["p"]
    llm = get_llm().with_structured_output(Gist)
    prompt = (
        f"You are a {p['name']} writing a short article gist for a market-research client, to FT / "
        f"Economist / WSJ standard. Decision: {p.get('decision', '')}. Headline: \"{p['title']}\". "
        f"The article's central model is {model_name}.\n\n"
        f"The executed model numbers you may cite (use the token in braces — NEVER type a digit):\n"
        f"{menu}\n\n"
        "Write ONE flowing paragraph of 150-200 words: the decisive read a strategist opens with, "
        "grounded in the model's insight and building to a clear call.\n"
        "Rules, non-negotiable:\n"
        "  • 150-200 words, one paragraph, decisive — no 'one reading', no hedging.\n"
        "  • Cite figures ONLY as the {n} tokens above, written exactly (e.g. {n0}). Never type, invent, "
        "round, or derive a number (no sums, ratios or differences of your own).\n"
        "  • Plain, confident desk-strategist voice; explain what the number MEANS, not just its value.\n"
        + (f"\nYour previous draft was rejected. Fix exactly this: {feedback}" if feedback else ""))
    return llm.invoke(prompt).body


def _finalize(body: str, toks: dict) -> tuple[str, str | None, int]:
    """Tokenise (convert any typed value matching an executed number → its {n}), leak-check the
    non-token remainder, then substitute verified values. Returns (filled_prose, leak_or_None, words)."""
    body_t = _tokenize(body, toks)
    stripped = _TOKLEFT.sub(" ", body_t)
    leak = _LEAK.search(stripped)
    filled = body_t
    for tok, no in toks.items():
        filled = filled.replace("{" + tok + "}", no.rendered())
    return filled, (leak.group() if leak else None), len(filled.split())


def _fill_template(mat: dict) -> str:
    """Deterministic fallback: the persona's authored summary_template, numbers injected from execution
    (the same fill as graph_corpus.build_persona_stub) — grounded, no LLM."""
    numbers = mat["numbers"]

    def _sub(m):
        k = m.group(1)
        if k in numbers:
            v = numbers[k].value
            return f"{v:.2f}" if isinstance(v, float) else str(v)
        return ""
    txt = re.sub(r"\{([^}]+)\}", _sub, mat["p"].get("summary_template", ""))
    return " ".join(txt.split())


def build_gist(mat: dict, model_id: str, model_name: str, *, max_iter: int = 3) -> tuple[str, str]:
    """Return (gist_text, source) where source ∈ {'llm','template'}. Firewall-guarded LLM prose, or the
    deterministic template on repeated failure — either way every figure is a verified executed value."""
    toks, menu = _citable_for_model(mat, model_id)
    if not toks:
        return _fill_template(mat), "template"
    feedback = ""
    for _ in range(max_iter):
        body = _narrate_gist(mat, model_name, menu, feedback=feedback)
        filled, leak, wc = _finalize(body, toks)
        if _TOKLEFT.search(filled):
            feedback = "you used a token that is not in the list; cite only the exact {n} tokens shown."
            continue
        if leak:
            feedback = (f"REMOVE the figure '{leak}' — it is not in the numbers list. Cite only the "
                        f"listed {{n}} tokens verbatim; never derive or invent a number.")
            continue
        if wc < 150 or wc > 205:
            feedback = f"your draft was {wc} words; write between 150 and 200 words."
            continue
        return filled, "llm"
    return _fill_template(mat), "template"


# ── docx assembly ────────────────────────────────────────────────────────────────────────────────
def _assemble_docx(path: Path, p: dict, mat: dict, model_name: str, ill_png: Path,
                   infog_png: Path | None, chart_paths: list[tuple[Path, str]], gist: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    doc.add_heading(p["title"], level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"{p['name']} — {p.get('decision', '')}").italic = True

    doc.add_picture(str(ill_png), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(gist)

    if infog_png and Path(infog_png).exists():
        doc.add_heading("At a glance", level=2)
        doc.add_picture(str(infog_png), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if chart_paths:
        doc.add_heading("The charts", level=2)
        for cp, cap in chart_paths:
            doc.add_picture(str(cp), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = doc.add_paragraph(cap)
            try:
                c.style = doc.styles["Caption"]
            except KeyError:
                pass
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    foot = doc.add_paragraph()
    fr = foot.add_run(f"Source: {', '.join(mat['source_labels']) or 'UMD'}. "
                      f"Data as of {mat['as_of']}. Model: {model_name}.")
    fr.font.size = Pt(8)
    fr.italic = True
    doc.save(str(path))


# ── the assembler ────────────────────────────────────────────────────────────────────────────────
def build_article(persona_id: str, conn, out_dir, *, backend: str = "auto", use_firewall: bool = True,
                  max_charts: int = 4, instance: str = "US") -> dict:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    mat = persona_material(persona_id, conn, instance=instance)
    p = mat["p"]
    model_id = MODEL_PICK[persona_id]
    run = mat["runs"][model_id]
    model_name = (run.get("meta") or {}).get("name", model_id)
    reasons: list[str] = []

    # 1. charts — the anchor model's, polished family render first, raw render as fallback
    chart_paths: list[tuple[Path, str]] = []
    for c in (run.get("charts") or []):
        b64, cap = chart_png_family(run, c["id"])
        if not b64:
            b64 = chart_png(run, c["id"])
            cap = decisive(" ".join((c.get("insight") or c["id"]).split()))
        if b64:
            cp = out_dir / f"chart_{len(chart_paths)}.png"
            cp.write_bytes(base64.b64decode(b64))
            chart_paths.append((cp, cap[:150]))
        if len(chart_paths) >= max_charts:
            break
    if not chart_paths:
        reasons.append("no charts rendered")

    # 2. infographic — the assigned family, decision_brief on any tier-1/eligibility failure
    infog_png = out_dir / "infographic.png"
    fam = FAMILY[persona_id]
    fam_name = fam.__name__.rsplit(".", 1)[-1]
    try:
        fam.render_persona(persona_id, conn, str(infog_png), instance=instance)
    except Exception as exc:
        reasons.append(f"{fam_name}→decision_brief: {str(exc).splitlines()[0][:80]}")
        fam_name = "decision_brief"
        try:
            decision_brief.render_persona(persona_id, conn, str(infog_png), instance=instance)
        except Exception as exc2:
            reasons.append(f"decision_brief failed: {str(exc2).splitlines()[0][:80]}")
            infog_png = None

    # 3. illustration — a van Gogh painting the art-director derives from THIS article's finding
    finding = _fill_template(mat)                           # the article's actual, number-filled content
    ill_b64, ill_meta = vangogh.illustration_png(
        finding, title=p["title"], decision=p.get("decision", ""),
        cache_key=f"{persona_id}|{model_id}|{instance}", backend=backend)
    ill_png = out_dir / "illustration.png"
    ill_png.write_bytes(base64.b64decode(ill_b64))

    # 4. gist — firewall-guarded LLM, template fallback
    gist, gist_src = (build_gist(mat, model_id, model_name) if use_firewall
                      else (finding, "template"))

    # 5. assemble
    docx_path = out_dir / "article.docx"
    _assemble_docx(docx_path, p, mat, model_name, ill_png, infog_png, chart_paths, gist)
    return {"persona": persona_id, "ok": True, "docx_path": str(docx_path),
            "caption": ill_meta.get("caption", ""), "scene": ill_meta.get("scene", ""),
            "family": fam_name, "model": model_id, "gist_src": gist_src, "gist": gist,
            "gist_words": len(gist.split()), "n_charts": len(chart_paths), "reasons": reasons,
            "illustration_png": str(ill_png), "infographic_png": (str(infog_png) if infog_png else ""),
            "chart_pngs": [str(cp) for cp, _ in chart_paths], "title": p["title"]}

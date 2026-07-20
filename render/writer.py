"""The full-article writer — a ~1000–1200-word feature, grounded in the models, anti-slop by construction.

Where `article.build_gist` writes one firewall-guarded paragraph, this writes the whole piece: a
standfirst that FORESHADOWS what the reader will learn, 4–6 body sections that weave the persona's whole
model set (each anchored to a chart), and a close that sits with the genuine open risk. Every number is
still authored by execution — the LLM cites values only by {token}; the renderer injects the verified
`NumberObject.rendered()`; any figure that matches no executed value is caught and rewritten.

The quality bar is FT / Economist / WSJ, and the explicit anti-pattern is StoryScope (arXiv 2604.03136):
AI prose over-explains and moralizes, tells tidy single-track stories, stays vague instead of naming
episodes, writes as if no one is watching, forces tidy resolution, over-writes the senses. Three layers
push against that: the StoryScope rules live in the prompts, a deterministic `_slop_lint` catches the
worst tells, and an LLM editor scored on the StoryScope checklist runs a bounded revise loop.
"""
from __future__ import annotations

import base64
import collections
import difflib
import os
import re
import sys
import textwrap
from functools import lru_cache
from pathlib import Path

import yaml
from pydantic import BaseModel, Field, model_validator

from .article import FAMILY, _fill_template
from .factsheet import sheets_for_run
from .illustration import vangogh
from .infographic.agentic import _citable, _tokenize
from .infographic.families import decision_brief
from . import graph_corpus
from .infographic.from_persona import GRAPH_DIR, chart_png, chart_png_family, decisive
from .judge.graph import judge_article
from .model_store import record_run
from .infographic.gate import _LEAK
from .studio.from_model import _refs
from .studio.llm import get_llm

_TOKLEFT = re.compile(r"\{n\d+\}")


def _invoke(llm, prompt: str, tries: int = 3):
    """Invoke a structured-output LLM, retrying on the intermittent parse failure where the model
    serializes a nested list as a string (a known .with_structured_output flake on long outputs)."""
    last = None
    for _ in range(tries):
        try:
            return llm.invoke(prompt)
        except Exception as exc:                            # ValidationError / parse error → re-ask
            last = exc
    raise last

# the worst AI-slop tells (StoryScope + the usual LLM crutches) — a deterministic pre-filter before the
# LLM critic. Widened from gate._HEDGE. Matched case-insensitively over the filled prose.
_SLOP = re.compile(
    r"\b(the (?:key )?takeaway|the lesson (?:here )?is|in conclusion|to sum up|it'?s worth noting|"
    r"at the end of the day|in a world where|one reading|it is important to note|needs? no introduction|"
    r"delve|tapestry|testament to|ever-evolving|navigat\w+ the complexit|in today'?s|"
    r"when it comes to|the bottom line|make no mistake|that said,? it'?s clear|"
    r"paints? a picture|speaks? volumes|the reality is|ultimately,|"
    # meta-scaffolding: narrating the article's own machinery instead of writing it (a clumsy AI tell —
    # address the reader about the SUBJECT, never describe 'this piece' or announce 'you will see')
    r"you will see|we will see|as we(?:'| wi)ll see|this piece|this article|the charts that follow|"
    r"the pivot of this|in the pages that follow|read on(?=[.,;:!?])|as this piece|the sections that follow)\b",
    re.I)


# ── schemas ──────────────────────────────────────────────────────────────────────────────────────
class SectionPlan(BaseModel):
    heading: str = Field(description="a short, specific section heading (not generic)")
    thesis: str = Field(description="the single point this section makes, in one sentence")
    model_id: str = Field(description="which model this section is grounded in")
    chart_ids: list[str] = Field(default_factory=list, description="0-2 chart ids (exact) that illustrate this section")
    token_ids: list[str] = Field(default_factory=list, description="the {n} tokens whose figures this section will cite")


class Outline(BaseModel):
    headline: str = Field(description="the article headline — keep or sharpen the given title; number-free")
    standfirst: str = Field(description="30-40 word standfirst that FORESHADOWS what the reader will learn "
                                        "without giving away the conclusion; may cite one {n} token")
    pivot: str = Field(description="the SINGLE mechanism the whole article turns on, to flag up front so "
                                   "the piece flows (e.g. 'the Gilchrist-Zakrajšek decomposition: is the "
                                   "spread paying you for default risk, or for risk appetite?')")
    sections: list[SectionPlan] = Field(description="4-6 sections forming a narrative arc that WEAVES the models")
    open_close: str = Field(description="the unresolved risk/tension the piece should end on — NOT a tidy bow")


class SectionDraft(BaseModel):
    heading: str
    prose: str = Field(description="the section body, flowing paragraphs; cite figures ONLY as {n} tokens")
    chart_ids: list[str] = Field(default_factory=list)


class SectionBinding(BaseModel):
    """A drafted section paired with the charts that belong to it — by identity, never by position.
    `model_id` is the planner's semantic anchor for the section; `origin` records how each chart got
    here (planned by the section editor / added by the exhibit contract to reach the floor / pulled in
    because the prose names it), for the human read and the completeness log."""
    heading: str
    model_id: str = ""
    prose: str = ""
    chart_ids: list[str] = Field(default_factory=list)
    origin: dict[str, str] = Field(default_factory=dict)   # chart_id -> planned | floor | prose


class Article(BaseModel):
    standfirst: str
    exec_summary: str = Field(description="the NARRATIVE executive summary — ~160-240 words of prose that "
                              "answers 'why should I read this?': an extended foreshadow that names the "
                              "article's central PIVOT up front, tells the reader what they will learn, and "
                              "DEFINES each key term/abbreviation once on first use (e.g. 'option-adjusted "
                              "spread, or OAS'). It sets up — and hands off to — the visual summary (the "
                              "infographic) that follows it. Cite figures only as {n} tokens.")
    sections: list[SectionDraft]


class Critique(BaseModel):
    ok: bool | None = None
    defects: list[str] = Field(default_factory=list,
                               description="concrete StoryScope/editorial failures: over-explaining or "
                               "moralizing, a tidy bow, vagueness where an episode should be named, no "
                               "foreshadow, purple/overwritten prose, generic filler, a weak lede")
    fixes: list[str] = Field(default_factory=list)

    @model_validator(mode="after")
    def _derive(self):
        if self.ok is None:
            self.ok = not self.defects
        return self


# ── the brief ────────────────────────────────────────────────────────────────────────────────────
def _active_says(meta: dict, latest) -> list[str]:
    """The interpretations whose `when` guard is TRUE on the latest run — the model's live regime call.

    The scope is INPUTS + OUTPUTS. It used to be outputs only, and guards legitimately reference
    inputs: reaction_function's `behind_the_curve` is `taylor_1993 - policy.level > 1.0`, where
    `policy` is an input. That raised NameError, the bare `except: continue` below swallowed it, and
    the interpretation had NEVER fired — while being TRUE (the gap is 1.68pp). An article headlined
    "how far behind the curve?" was never told by its own model that the Fed is behind the curve.
    Inputs are §10 State objects, so `policy.level` resolves once they are in scope; outputs win a
    name collision, being the model's product.

    A guard that RAISES is now loud. A silent failure here is indistinguishable from "the condition
    is false", which is exactly how this hid: 34 of 46 guards are legitimately false today, so one
    that could never fire looked like all the others.
    """
    if latest is None:
        return []
    scope = {**(getattr(latest, "inputs", {}) or {}), **(getattr(latest, "outputs", {}) or {})}
    says = []
    for interp in (meta.get("interpretations") or []):
        when = interp.get("when")
        if not when:
            continue
        try:
            if eval(when, {"__builtins__": {}}, scope):    # noqa: S307 — trusted catalog exprs, no builtins
                says.append(" ".join(str(interp.get("says", "")).split()))
        except Exception as exc:
            print(f"INTERPRETATION BROKEN — {meta.get('model_id')}/{interp.get('id')}: "
                  f"`{when}` raised {type(exc).__name__}: {exc}", file=sys.stderr)
    return says


def _is_snapshot(c: dict) -> bool:
    """A static 'today' bar / named-values chart — redundant when a time series of the same data exists
    (the series shows the snapshot AND its change; the bar only the snapshot)."""
    dc = c.get("data_contract", {}) or {}
    return dc.get("kind") in {"named_values", "bar"} or c.get("chart_type") in {"bar", "named_values"}


def build_brief(mat: dict, *, conn=None, limit: int = 24) -> dict:
    """Everything the planner/writer sees: the token menu, and per model its name, grounding, method,
    outputs, chart insights, live regime call, and — when `conn` is given — the §10 FACT SHEET of every
    output series it is asked to narrate. Plus a chart index for id resolution.

    The fact sheet is the difference between narrating FROM the model and narrating ABOUT it. Without
    it the Narrator knows two strings per chart — a title and a catalog `insight` written months ago —
    so it recites the textbook's story rather than this run's. That is why "r* drifted down through the
    2010s" and "the most restrictive setting since the financial crisis" shipped: not carelessness, but
    the only thing possible for a writer that has never seen the series.

    Recording the runs here is deliberate. The Narrator and the Judge must read the SAME source, or the
    check is against a different set of numbers than the prose was written from. `runs` is returned so
    the Judge can adjudicate against exactly these rows.
    """
    toks, menu = _citable(mat, limit=limit)
    models, chart_index, runs = [], {}, {}
    for mid in mat["p"].get("models", []):
        run = mat["runs"].get(mid) or {}
        meta = run.get("meta") or {}
        if not meta:
            continue
        sheet = ""
        if conn is not None:
            # A fact sheet is evidence; failing to produce one must be visible, never a silent
            # downgrade to the priors-reciting Narrator we are trying to retire.
            try:
                rid = record_run(conn, run, instance=run.get("instance"))
                if rid:
                    runs[mid] = rid
                    sheet = sheets_for_run(conn, rid, meta.get("outputs") or [])
            except Exception as exc:
                print(f"FACT SHEET UNAVAILABLE — {mid}: {type(exc).__name__}: {exc}", file=sys.stderr)
        sp = meta.get("spec") or {}
        equations = sp.get("equations", "") if isinstance(sp, dict) else ""
        raw = run.get("charts") or []
        rich = [c for c in raw if not _is_snapshot(c)]
        keep = rich if rich else raw                       # keep the static snapshot only if nothing richer
        charts = []
        for c in keep:
            cid, ins = c.get("id", ""), " ".join((c.get("insight") or "").split())
            if cid:
                charts.append((cid, ins))
                # `role` (input/outcome/consequence) and `refs` feed the exhibit contract: role drives
                # diversity when filling to the floor, refs identify near-duplicates. Both were already
                # in the catalog and neither was read.
                chart_index[cid] = {
                    "model_id": mid, "insight": ins, "role": c.get("role"),
                    "refs": frozenset(_refs(c.get("data_contract") or {})),
                    "kind": (c.get("data_contract") or {}).get("kind", ""),
                }
        outs = "; ".join(f"{o['name']} ({o.get('unit','')}) — {o.get('meaning','')}"
                         for o in (meta.get("outputs") or []))
        models.append({
            "id": mid, "name": meta.get("name", mid),
            "grounded_in": ", ".join(meta.get("grounded_in") or []),
            "method": meta.get("method_note", "") or equations,
            "outputs": outs, "charts": charts,
            "regime": _active_says(meta, run.get("latest")),
            "sheet": sheet,
        })
    # The data boundary: the earliest observation any of this persona's models actually holds. The writer
    # must NOT name a market episode before this (the data cannot show it) — the firewall for HISTORY.
    starts = [str((run.get("history") or [{}])[0].as_of)[:10]
              for mid in mat["p"].get("models", [])
              if (run := mat["runs"].get(mid)) and run.get("history")]
    data_start = min(starts) if starts else ""
    return {"mat": mat, "toks": toks, "menu": menu, "models": models,
            "chart_index": chart_index, "papers": mat.get("papers", []),
            "as_of": mat.get("as_of", ""), "data_start": data_start, "runs": runs}


def _brief_text(brief: dict) -> str:
    lines = ["THE EVIDENCE (every figure is executed on real data; cite ONLY via the {n} tokens):", ""]
    ds = brief.get("data_start", "")
    if ds:
        yr = ds[:4]
        lines.append(f"⚠ DATA WINDOW: every chart begins {ds}. You may name ONLY market episodes that fall "
                     f"on or after {yr} (e.g. events the data actually shows). You MUST NOT reference any "
                     f"episode before {yr} — the 2008 GFC, 2011, the taper tantrum, etc. are NOT in this "
                     f"data and claiming a chart 'prints' them is a fabrication. Name only what the data spans.")
        lines.append("")
    lines.append("Citable figures:")
    lines.append(brief["menu"])
    lines.append("")
    for m in brief["models"]:
        lines.append(f"■ MODEL: {m['name']}  [grounded in: {m['grounded_in'] or 'n/a'}]")
        if m["method"]:
            lines.append(f"   method: {m['method']}")
        if m["outputs"]:
            lines.append(f"   outputs: {m['outputs']}")
        if m.get("sheet"):
            lines.append("   ► WHAT THIS MODEL ACTUALLY PRODUCED IN THIS RUN — read it before you write a "
                         "word about shape. Any claim about a HIGH/LOW ('the most X since Y'), a SIGN "
                         "('below zero', 'inverted') or a MOVE ('has risen', 'more recently up') must "
                         "match these numbers. They are executed, they are this run's, and they OVERRIDE "
                         "anything you remember about how this series is supposed to behave:")
            lines.append(m["sheet"])
        for cid, ins in m["charts"]:
            lines.append(f"   chart «{cid}» — {ins}")
        for s in m["regime"]:
            lines.append(f"   ► LIVE REGIME CALL: {s}")
        lines.append("")
    return "\n".join(lines)


# ── the StoryScope voice (shared by planner + writer) ──────────────────────────────────────────────
_VOICE = (
    "Write to the standard of the Financial Times, The Economist or the Wall Street Journal — proper "
    "journalism, not AI filler. Heed these hard rules (they are the difference between a real piece and "
    "generated slop):\n"
    "  • INFER, don't explain. Never state the moral or 'the takeaway'; never tell the reader what to "
    "think. Lay out the evidence and trust the reader to draw the conclusion.\n"
    "  • FORESHADOW then delay. Open with a hook that promises what the reader will understand by the "
    "end; do not resolve it in the first paragraph. Build tension across the piece.\n"
    "  • Name the SPECIFIC episodes THE DATA ACTUALLY SPANS (see the DATA WINDOW above) — specificity is "
    "the mark of a human writer, vague allusion the mark of a machine. But NEVER name an episode that "
    "predates the data window: if the series starts in 2016, the 2008 crisis is not yours to cite.\n"
    "  • Address the reader where it earns its place ('you are being paid…'). Do not write as if no one "
    "is watching.\n"
    "  • End on the genuine unresolved risk, NOT a tidy bow. Sit with the ambiguity.\n"
    "  • Plain and concrete over ornate. No purple metaphor, no bodily/sensory overwriting, no throat-"
    "clearing ('it is worth noting', 'in today's market', 'make no mistake').\n"
    "  • DIGITS MEAN MEASURED. WORDS MEAN APPROXIMATE. This is a hard contract and it runs both ways.\n"
    "      – Every MEASURED value is cited ONLY as the {n} token provided, written exactly. Never type "
    "a digit of your own. Each token ALREADY carries its unit (it may render as '4.55 vol pts' or "
    "'14%') — write the token alone and NEVER restate the unit after it (never '{n} vol points', "
    "never '{n} percent').\n"
    "      – ALWAYS CONNECT A FIGURE TO THE SENTENCE with a preposition or punctuation — 'at', 'of', "
    "'to', 'near', 'by', or a comma. NEVER jam a bare descriptor straight into the token: write 'sits "
    "at a loose {n}', 'positive, at {n}', 'the premium is only {n} of the yield' — NEVER 'sits loose "
    "{n}', 'positive {n}', 'fifteen percent risk premium {n}'. A word run into a figure with no "
    "connector reads as unfinished; it is the single most common blemish, so guard every token.\n"
    "      – When the prose genuinely needs a number FIGURATIVELY — a rough magnitude, a round "
    "threshold, a target, a rule of thumb — SPELL IT OUT IN WORDS and never in digits: 'more than a "
    "hundred basis points below its own rule', 'an r* of around one percent', 'the two percent "
    "target', 'peaks near twenty'. This is allowed, it is good FT/Economist style, and it is how you "
    "say 'roughly' without pretending to a precision you do not have. The spelled-out words REPLACE "
    "the figure — NEVER place them next to the token or numeral for the same number (never 'above "
    "five percent {n}', never 'around one percent 1.06%').\n"
    "      – NEVER spell out a PRECISE value to dodge a token. 'two point three nine percent' is a "
    "measured figure in disguise and is the one thing you must not write. If you know the exact "
    "number, it has a token — use the token.\n"
    "    The reader must be able to tell at a glance which numbers are measured and which are "
    "gestures. Digits are a promise that the number came from the model.\n"
    "  • Use real terms: never a garbled or invented phrase (write 'at-the-money', never a mangled "
    "substitute).")


def plan_arc(brief: dict, feedback: str = "") -> Outline:
    p = brief["mat"]["p"]
    llm = get_llm(max_tokens=4096).with_structured_output(Outline)
    prompt = (
        f"You are the section editor planning a ~1000-1200 word feature for a {p['name']}. "
        f"The decision it informs: {p.get('decision','')}. Working title: \"{p['title']}\".\n\n"
        f"{_brief_text(brief)}\n\n"
        "First decide the PIVOT — the single mechanism the whole piece turns on — and plan for it to be "
        "flagged UP FRONT (in the standfirst and the executive summary) so the article flows: tell the "
        "reader what you are going to tell them, then tell them. Then plan the article as a NARRATIVE ARC "
        "that WEAVES these models into one thesis — foreshadow, develop across 4-6 sections, turn on the "
        "sharpest finding, end on the open risk. Do NOT write one tidy section per model in list order; "
        "braid them. Assign each section the exact chart id(s) that illustrate it and the {n} tokens it "
        "will cite. CHART CHOICE MATTERS: prefer the chart that shows the most. When a point is about a "
        "CROSS-SECTION of levels (a curve, a quality ladder, a term structure), place the chart that shows "
        "BOTH the levels AND how they have moved — a now-vs-prior comparison ('now vs 3 months ago') or "
        "the cross-section through time — in preference to a lone derived spread or a single static "
        "snapshot; a snapshot that only shows today's levels is redundant with such a chart and must not "
        "displace it. One strong chart per point.\n\n" + _VOICE
        + (f"\n\nYour previous plan was rejected. Fix: {feedback}" if feedback else ""))
    return _invoke(llm, prompt)


def write_article(brief: dict, outline: Outline, feedback: str = "") -> Article:
    p = brief["mat"]["p"]
    llm = get_llm(max_tokens=8192).with_structured_output(Article)
    plan = "\n".join(
        [f"HEADLINE: {outline.headline}", f"STANDFIRST (foreshadow): {outline.standfirst}",
         f"PIVOT (flag this up front): {outline.pivot}",
         f"END ON (open risk): {outline.open_close}", "", "SECTIONS:"]
        + [f"  {i+1}. «{s.heading}» — {s.thesis}  [charts: {', '.join(s.chart_ids) or 'none'}; "
           f"cite: {', '.join(s.token_ids) or 'none'}]" for i, s in enumerate(outline.sections)])
    prompt = (
        f"You are a {p['name']} writing the full feature to the plan below. Target 1000-1200 words total.\n"
        "STRUCTURE, in order:\n"
        "  1. STANDFIRST — the sharp foreshadowing hook (given).\n"
        "  2. EXECUTIVE SUMMARY — ~160-240 words of narrative prose that earns the reader's next ten "
        "minutes. Establish the stakes and the PIVOT, and DEFINE each key term/abbreviation once on first "
        "use (e.g. 'the option-adjusted spread, or OAS', 'high yield (HY)'), using the abbreviation "
        "thereafter. Foreshadow with FLAIR, the way an FT or Economist nut-graf does — through a confident, "
        "specific claim and the tension in it, so the promise of what's coming is IMPLICIT. Do NOT recite a "
        "table of contents. FORBIDDEN, because they are the clumsy mechanical tell an editor would cut: "
        "'you will see…', 'this piece/article', 'the pivot of this piece', 'the charts that follow', "
        "'start with the…', and any sentence that narrates the article's own structure. Address the reader "
        "about the SUBJECT ('you are being paid…'), never about the document.\n"
        "  3. SECTIONS — the detail: flowing paragraphs (not bullets), each grounded in its model and "
        "interpreting its chart(s) in words. INTERPRET every chart you point at — say what the reader "
        "should see in it and why it matters; never name a chart without reading it.\n\n"
        f"{_brief_text(brief)}\n\n{plan}\n\n" + _VOICE
        + "\n\nReturn the standfirst, the executive summary, and each section's heading + prose. Cite "
        "figures ONLY as the {n} number tokens — the ONLY braces permitted in your prose. Do NOT insert "
        "figure numbers, footnote markers, or {chart} references; to point at a chart, name it in plain "
        "words. The charts are placed for you from the plan."
        + (f"\n\nYour previous draft was rejected. Fix exactly this: {feedback}" if feedback else ""))
    return _invoke(llm, prompt)


# ── firewall + slop lint ───────────────────────────────────────────────────────────────────────────
_ONES = (r"zero|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|"
         r"fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|"
         r"eighty|ninety|hundred")
# "two point three nine percent" — a MEASURED figure wearing words to slip past a digit-matching
# firewall. The writer may spell out approximations ("around one percent", "more than a hundred basis
# points"); that is the deal, and it is what lets the prose be figurative without inventing data. This
# is the other half of the deal. It matters more now than it used to: since the fact sheet landed the
# narrator can SEE the real numbers, so it has something precise to disguise.
# Trailing digit-words are consumed so the whole figure is quoted back ("two point three nine", not
# "two point three") — the feedback names the string, and half a number reads like a typo.
_WORD_DECIMAL = re.compile(rf"\b(?:{_ONES})[\s-]+point[\s-]+(?:{_ONES})(?:[\s-]+(?:{_ONES}))*\b", re.I)


def _worded_precision(text: str) -> str | None:
    m = _WORD_DECIMAL.search(text)
    return m.group() if m else None


# ── "one figure, once" (number-rendering guard) ──────────────────────────────────────────────────
# A figure must appear ONCE. Two ways it doubled in v3, both invisible to the claim-checker because
# they are typography, not claims: (a) a token cited twice — "reads +0.18pp +0.18pp"; (b) a spelled-out
# round approximation run into its own numeral token — "above five percent 5.31%", "about one percent
# 1.06%". (b) is SELF-INFLICTED: the leak-feedback tells the writer to spell out round thresholds to
# clear the digit firewall, and the token for that same figure then expands right beside the words.
_QUAL = (r"about|around|above|below|just|only|nearly|roughly|over|under|almost|barely|some|"
         r"north of|south of|more than|less than|at least|at most")
_ADJ_DUP = re.compile(r"([+\-−]?\d[\d,]*(?:\.\d+)?\s?(?:%|pp|bp|bps|×|σ))\s+\1(?!\w)")
_WORDED_ROUND_PCT = re.compile(
    rf"\b((?:(?:{_QUAL})\s+)*)(?:a\s+)?(?:{_ONES})(?:[\s-]+(?:{_ONES}))*\s+per\s?cent\s+"
    rf"([+\-]?\d[\d,]*(?:\.\d+)?\s?%)", re.I)
_WORDED_ROUND_BP = re.compile(
    rf"\b((?:(?:{_QUAL})\s+)*)(?:a\s+)?(?:{_ONES})(?:[\s-]+(?:{_ONES}))*\s+basis\s+points?\s+"
    rf"([+\-]?\d[\d,]*(?:\.\d+)?\s?(?:bp|bps))", re.I)

# v4 (N1) added two forms my adjacent guard missed. (a) a NOUN PHRASE between the spelled figure and its
# numeral — "only fifteen percent risk premium 15%" (the 10-year deck). (b) a bare state DESCRIPTOR run
# straight into a signed/σ token with no connector — "sits loose -0.52σ", "positive +0.35σ" (pervasive
# in fin-conditions). The durable fix for (b) is the writer prompt; this is the backstop.
_NUMWORD = {"zero": 0, "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6, "seven": 7,
            "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12, "thirteen": 13, "fourteen": 14,
            "fifteen": 15, "sixteen": 16, "seventeen": 17, "eighteen": 18, "nineteen": 19, "twenty": 20,
            "thirty": 30, "forty": 40, "fifty": 50, "sixty": 60, "seventy": 70, "eighty": 80, "ninety": 90,
            "hundred": 100}
_WORDED_PHRASE_NUM = re.compile(
    rf"\b((?:{_ONES})(?:[\s-]+(?:{_ONES}))*)\s+per\s?cent\s+((?:[a-z][a-z-]+\s+){{1,3}}?)"
    rf"(\d+(?:\.\d+)?)\s?%", re.I)
_STATE_ADJ = (r"loose|tight|elevated|depressed|positive|negative|high|low|firm|soft|hot|cold|cheap|"
              r"rich|wide|narrow|steep|flat|inverted|extended|stretched|compressed|subdued|muted|benign")
_ADJ_NUM = re.compile(rf"\b({_STATE_ADJ})\s+([+\-−][\d.]+\s?(?:σ|pp|bp|bps|%)|[\d.]+\s?σ)", re.I)


def _spelled_value(phrase: str):
    tot = cur = 0
    for w in re.findall(r"[a-z]+", phrase.lower()):
        if w not in _NUMWORD:
            continue
        v = _NUMWORD[w]
        if v == 100:
            cur = (cur or 1) * 100
        else:
            cur += v
    return (tot + cur) or None


def _one_figure_once(t: str, toks: dict) -> str:
    """Collapse a figure that renders twice, and connect a descriptor jammed into a numeral. A figure
    appears ONCE, with a connector — never doubled, never a bare word run into its own token."""
    for no in toks.values():                                   # exact adjacent dup (incl. word units)
        v = _render_for_prose(no)
        if v:
            t = re.sub(rf"({re.escape(v)})\s+{re.escape(v)}(?!\w)", r"\1", t)
    t = _ADJ_DUP.sub(r"\1", t)                                 # any adjacent identical figure
    t = _WORDED_ROUND_PCT.sub(lambda m: (m.group(1) or "") + m.group(2), t)
    t = _WORDED_ROUND_BP.sub(lambda m: (m.group(1) or "") + m.group(2), t)

    def _phrase(m):                                            # "fifteen percent risk premium 15%"
        sv = _spelled_value(m.group(1))
        if sv is not None and abs(sv - float(m.group(3))) < 0.5:   # same figure, spelled AND tokenised
            return f"{m.group(1)} percent {m.group(2).strip()}"    # keep the spelled phrase, drop the numeral
        return m.group(0)
    t = _WORDED_PHRASE_NUM.sub(_phrase, t)
    t = _ADJ_NUM.sub(lambda m: f"{m.group(1)}, at {m.group(2)}", t)   # "loose -0.52σ" → "loose, at -0.52σ"
    return t


def _render_for_prose(no) -> str:
    """A NumberObject's value for running prose: like rendered(), but a multi-word unit ('vol pts') gets
    a space off the number ('4.55 vol pts', not '4.55vol pts'); tight units (%, pp, bp, σ) stay glued."""
    r = no.rendered()
    u = (no.unit or "").strip()
    if u and " " in u and r.endswith(u):
        return r[: -len(u)].rstrip() + " " + u
    return r


def _finalize_text(text: str, toks: dict) -> tuple[str, str | None]:
    """Tokenise any typed value that matches an executed number, leak-check the non-token remainder, then
    substitute verified values. Returns (filled_text, leak_or_None)."""
    t = _tokenize(text, toks)
    stripped = _TOKLEFT.sub(" ", t)
    leak = _LEAK.search(stripped)
    for tok, no in toks.items():
        val = _render_for_prose(no)
        # SIGN DOUBLING. The writer does not know the token already carries its sign, so it writes
        # "+{n}" and the reader gets "a stance of just ++0.18pp" — which shipped on 2026-07-17, in an
        # article the judge passed as fully grounded. The arithmetic was right and the typography was
        # garbage, and no claim-checker will ever catch that: it is not a claim.
        # Only a sign that AGREES with the rendered one is absorbed. A mismatch ("-{n}" where n is
        # positive) means the writer meant something else, and quietly rewriting it would hide a real
        # error behind a tidy-looking number.
        if val[:1] in "+-":
            t = re.sub((r"\+\s*" if val[0] == "+" else r"[-−–]\s*") + r"\{" + re.escape(tok) + r"\}",
                       lambda _m, v=val: v, t)
        t = t.replace("{" + tok + "}", val)
    # belt-and-suspenders: strip a unit the writer restated right after a token expansion, e.g.
    # "4.55 vol pts vol points" -> "4.55 vol pts" (the prompt tells it not to, this catches slips).
    for no in toks.values():
        u = (no.unit or "").strip()
        if u and " " in u:                                  # word unit like "vol pts"
            head = re.escape(u.split()[0])                  # "vol"
            t = re.sub(rf"({re.escape(u)})\s+{head}\s+\w+", r"\1", t, flags=re.I)
    t = _one_figure_once(t, toks)                           # a figure must appear once, not twice
    return t, (leak.group() if leak else None)


def _finalize_article(article: Article, toks: dict) -> tuple[str, str, str, list[SectionDraft], str | None]:
    """Fill the standfirst + exec summary + every section; return
    (full_text, standfirst, exec_summary, filled_sections, first_leak)."""
    leaks: list[str] = []
    sf, lk = _finalize_text(article.standfirst, toks)
    if lk:
        leaks.append(lk)
    ex, lk = _finalize_text(article.exec_summary, toks)
    if lk:
        leaks.append(lk)
    filled_sections = []
    for s in article.sections:
        prose, lk = _finalize_text(s.prose, toks)
        if lk:
            leaks.append(lk)
        filled_sections.append(SectionDraft(heading=s.heading, prose=prose, chart_ids=s.chart_ids))
    full = sf + "\n\n" + ex + "\n\n" + "\n\n".join(f"{s.heading}\n{s.prose}" for s in filled_sections)
    unfilled = _TOKLEFT.search(full)
    leak = (unfilled.group() if unfilled else None) or (leaks[0] if leaks else None)
    return full, sf, ex, filled_sections, leak


def _slop_lint(text: str) -> list[str]:
    return sorted({m.group(0) for m in _SLOP.finditer(text)})


_YEAR = re.compile(r"\b(199\d|20\d\d)\b")   # market-episode years 1990–2099 (pre-1990 = academic cites, ignored)


def _episode_leaks(text: str, data_start: str) -> list[str]:
    """Deterministic HISTORY firewall: any year named in the prose that predates the data window is a
    fabricated episode (the charts cannot show it). Returns the offending years. The 1990 floor lets
    genuine pre-1990 academic citations (GARCH 1986, etc.) through — market crises are all post-1990."""
    if not data_start:
        return []
    start = int(data_start[:4])
    bad = []
    for m in _YEAR.finditer(text):
        y = int(m.group(0))
        if not (1990 <= y < start):
            continue
        # skip a lone parenthesised year — that's an academic citation ("Taylor (1993)"), not a market
        # episode the chart claims to show. A crisis LIST reads "(2008, 2011, …)" — comma after → flagged.
        if text[m.start() - 1:m.start()] == "(" and text[m.end():m.end() + 1] == ")":
            continue
        bad.append(y)
    return [str(y) for y in sorted(set(bad))]


# ── further reading: resolve grounding slugs → real citations (title + summary + url) ───────────────
_REG_PATH = Path(__file__).resolve().parents[1] / "knowledge" / "registry.yaml"


@lru_cache(maxsize=1)
def _registry() -> dict:
    try:
        return {p["id"]: p for p in yaml.safe_load(_REG_PATH.read_text()).get("papers", [])}
    except Exception:
        return {}


def _resolve_papers(slugs: list[str]) -> list[dict]:
    """Grounding slugs → reader-facing citations {title, byline, summary, url} from the corpus registry."""
    reg, out = _registry(), []
    for s in slugs:
        p = reg.get(s)
        if not p:
            continue
        authors = p.get("authors") or []
        byline = (", ".join(authors[:2]) + (" et al." if len(authors) > 2 else "")) if authors else ""
        why = " ".join(str(p.get("why", "")).split())
        summary = re.split(r"(?<=[.;])\s", why)[0] if why else ""       # first sentence only
        out.append({"title": p.get("title", s), "byline": byline, "summary": summary,
                    "url": p.get("url", "")})
    return out


def _add_hyperlink(paragraph, url: str, text: str):
    """A real clickable hyperlink run (python-docx has no native helper)."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement
    r_id = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    und = OxmlElement("w:u"); und.set(qn("w:val"), "single"); rpr.append(und)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); paragraph._p.append(link)


@lru_cache(maxsize=1)
def _byline() -> dict:
    """The author identity (catalog/byline.yaml) for the byline + closing 'work with me' block. Blank
    fields are omitted so an unfinished byline never ships; missing file → no hook at all."""
    try:
        d = yaml.safe_load((GRAPH_DIR.parent / "byline.yaml").read_text()) or {}
    except Exception:
        return {}
    return {k: (str(v).strip() if v is not None else "") for k, v in d.items()}


def _strip_stray(text: str) -> str:
    """Final safety net: remove any leftover {…} marker (e.g. an invented figure ref) and tidy spacing —
    so a broken placeholder can never ship even if the writer keeps re-inserting them."""
    text = re.sub(r"\s*\{[^}]*\}", "", text)
    return re.sub(r"[ \t]{2,}", " ", re.sub(r"\s+([.,;:])", r"\1", text)).strip()


def _judge(full_text: str, brief: dict, conn) -> tuple[bool, list]:
    """§06 role 7 — is the prose grounded in what the models actually produced?

    Returns (grounded, failures). A Judge that cannot run must not report `grounded=True`: silence
    passing for a pass is the exact failure this role exists to end (the old critic recorded
    `critic_ok=True` beside "the most restrictive since the financial crisis"). If it breaks, say so
    and treat the article as unproven rather than clean.
    """
    runs = brief.get("runs") or {}
    if not runs or conn is None:
        print("JUDGE SKIPPED — no recorded runs to adjudicate against; grounding is UNPROVEN",
              file=sys.stderr)
        return False, []
    try:
        st = judge_article(full_text, runs, conn)
        return bool(st.get("grounded")), list(st.get("failures") or [])
    except Exception as exc:
        print(f"JUDGE FAILED — {type(exc).__name__}: {exc}; grounding is UNPROVEN", file=sys.stderr)
        return False, []


def critique_article(full_text: str, title: str) -> Critique:
    llm = get_llm(max_tokens=2048).with_structured_output(Critique)
    prompt = (
        f"You are a demanding FT/Economist editor reviewing this feature draft (headline: \"{title}\"). "
        "Judge it ONLY against the marks of real journalism vs AI slop. FAIL it for any of: (1) it "
        "over-explains or moralizes / states 'the takeaway' / tells the reader what to think; (2) it "
        "resolves into a tidy bow instead of ending on a real open risk; (3) it is vague where it should "
        "name a specific episode the data actually spans; (4) the standfirst does not foreshadow, or the lede "
        "resolves everything at once; (5) purple or overwritten prose, bodily/sensory padding; (6) generic "
        "filler, throat-clearing, or a section that merely restates a number without a point; (7) the "
        "foreshadowing is PROSAIC or MECHANICAL — a table-of-contents recital ('you will see X, you will "
        "see Y'), or it narrates the article's own structure ('this piece', 'the charts that follow'), "
        "instead of the confident, implicit nut-graf foreshadowing a top paper would run; (8) any "
        "GARBLED or non-English phrase, a mangled term ('at-metric-of-the-moment' for 'at-the-money'), "
        "or a unit restated after a figure ('4.55 vol pts vol points'). If it reads "
        "like a real desk strategist wrote it — specific, confident, arced, honest about the open risk — "
        "set ok=true. Give concrete fixes.\n\nDRAFT:\n" + full_text[:9000])
    return _invoke(llm, prompt)


# ── docx assembly ──────────────────────────────────────────────────────────────────────────────────
# ── the exhibit contract ─────────────────────────────────────────────────────────────────────────
# The charts ARE the argument — every section is a chart read aloud — so the exhibit count is not the
# LLM's to decide. It was, and it swung wildly: the 2026-07-14 run shipped 10 charts (reviewers: "twelve
# exhibits is two or three too many"); the 2026-07-16 run shipped ONE, from the same pipeline, with the
# same 13 charts offered to the planner. _place_charts's only safety net fired at ZERO
# (`if sum(...) == 0`), so a single chart sailed through as though it were a choice.
#
# So: the planner still chooses WHICH charts (it saw the index; that judgement is worth having). This
# decides HOW MANY. LLM proposes, arithmetic disposes — the same rule as every other gate here.
#
# 4-6 charts + the illustration + the infographic = 6-8 exhibits, which is what the editorial reviews
# asked for ("the target titles would run six to eight").
_CHART_FLOOR, _CHART_CEILING = 4, 6
# No more than this many charts from one model. "Four charts for the funding split" was the
# treasurer review's complaint; the credit piece showed the EBP twice. A persona with 2 models can
# still reach the floor (2x2=4); the cap RELAXES rather than let an article ship under the floor,
# because too few charts is the worse failure.
_MAX_PER_MODEL = 2


def _chart_key(cid: str, ci: dict) -> tuple:
    """Near-duplicate identity: the same model drawing the same refs is the same exhibit twice.
    This is the `redundant-exhibits` defect (6/8 reviews) — the credit piece showed the EBP twice, the
    treasurer four views of one decomposition — and the signal was always here, unused."""
    e = ci.get(cid) or {}
    return (e.get("model_id"), e.get("refs"))


_LOCUS_WORDS = ("scatter", "locus", "against", "outward", "shift", "corner", "top-left",
                "bottom-right", "quadrant", "cloud")
_SERIES_WORDS = ("over time", "through time", "two series", "history", "by tenor", "each leg", "legs")
_GAP_WORDS = ("gap", "minus", "difference", "consequence", "how far", "the trade", "below its rule",
              "above its rule", "shortfall", "distance from")
_GAP_IDS = ("gap", "consequence", "minus", "how far", "difference", "shortfall")


def _prefer_form(a: str, b: str, ci: dict, prose_l: str) -> str:
    """Of two same-concept charts, keep the one the section's prose FORM cues point at (a locus/scatter
    when the prose describes a shift/corner, a series when it describes movement over time); otherwise
    the richer (more referenced fields). This is why the labour Okun SCATTER survives over its two-line
    twin — the corrected prose describes the locus."""
    ea, eb = ci.get(a) or {}, ci.get(b) or {}
    a_rel = (ea.get("kind") or "").lower() in ("scatter", "pearson") or \
        any(w in a.lower() for w in ("vs", "versus", "against"))
    b_rel = (eb.get("kind") or "").lower() in ("scatter", "pearson") or \
        any(w in b.lower() for w in ("vs", "versus", "against"))
    if a_rel != b_rel and any(w in prose_l for w in _LOCUS_WORDS):
        return a if a_rel else b
    if a_rel != b_rel and any(w in prose_l for w in _SERIES_WORDS):
        return b if a_rel else a
    # gap/difference vs a levels chart of the same fields: keep the GAP when the prose is about the gap
    # (macro_rates called "the consequence chart — prescription minus actual — the one to sit with", and
    # the old "richer = more series" tie-break dropped it for the levels fan).
    a_gap = (ea.get("kind") or "").lower() == "gap_series" or any(w in a.lower() for w in _GAP_IDS)
    b_gap = (eb.get("kind") or "").lower() == "gap_series" or any(w in b.lower() for w in _GAP_IDS)
    if a_gap != b_gap and any(w in prose_l for w in _GAP_WORDS):
        return a if a_gap else b
    return a if len(ea.get("refs") or ()) >= len(eb.get("refs") or ()) else b


def _dedup_within(cids: list[str], ci: dict, prose: str) -> list[str]:
    """Within ONE section, collapse a subset chain of same-concept charts to a single exhibit — the
    double-EBP, or a locus and its two-series twin — keeping the prose-described form."""
    prose_l = (prose or "").lower()
    out: list[str] = []
    for cid in cids:
        refs = (ci.get(cid) or {}).get("refs") or frozenset()
        if not refs:
            out.append(cid)
            continue
        drop_self, keep = False, []
        for kc in out:
            kr = (ci.get(kc) or {}).get("refs") or frozenset()
            if refs <= kr or kr <= refs:                  # same concept (one set contains the other)
                if _prefer_form(cid, kc, ci, prose_l) == kc:
                    drop_self = True
                    keep.append(kc)
                # else drop kc (omit) — cid is the kept form, appended below
            else:
                keep.append(kc)
        out = keep
        if not drop_self:
            out.append(cid)
    return out


def _dedup_bindings(bindings: list[SectionBinding], ci: dict) -> list[str]:
    """Per-section subset-collapse (prose-form-aware), then across sections drop only near-EXACT
    duplicates. NEVER a cross-section subset: a gap chart in one section and its parent levels-fan in
    another are BOTH wanted (the Fed regression). Returns the ids dropped, for logging."""
    before = [c for b in bindings for c in b.chart_ids]
    for b in bindings:
        b.chart_ids = _dedup_within(b.chart_ids, ci, b.prose)
    # cross-section decomposition-collapse runs BEFORE the exact-dup pass so the survivor is chosen by
    # which section ARGUES the concept, not by iteration order (an exact twin would else keep the opener).
    _dedup_decomposition_cross_section(bindings, ci)
    seen: dict = {}
    for b in bindings:
        keep = []
        for c in b.chart_ids:
            k = _chart_key(c, ci)
            if k in seen:                                 # exact/near duplicate already shown earlier
                continue
            seen[k] = c
            keep.append(c)
        b.chart_ids = keep
    kept = {c for b in bindings for c in b.chart_ids}
    return [c for c in before if c not in kept]


# Over-time decomposition FORMS only. A gap chart (`gap_series`) and its parent levels fan (`series`) in
# two sections are BOTH wanted (the Fed regression) — they are not in this set, so they are never touched.
_DECOMP_FORMS = ("decomposition", "stacked")


def _dedup_decomposition_cross_section(bindings: list[SectionBinding], ci: dict) -> None:
    """The double-EBP re-grows when Role-2 imports a model whose over-time decomposition already lives in
    another section: the SAME model's spread-decomposition shown twice, once in the opener and once where
    it's argued (6/8 reviews flagged it). Collapse those to ONE — kept in the section whose prose best
    reads it — while leaving a snapshot bar, a gap chart, or a different model's decomposition alone.

    Restricted to over-time decomposition forms of the SAME model so it can only ever fire on a genuine
    same-concept duplicate, never on the gap/levels pair the per-section pass is explicitly told to keep."""
    by_model: dict = {}
    for b in bindings:
        for c in b.chart_ids:
            e = ci.get(c) or {}
            if (e.get("kind") or "").lower() in _DECOMP_FORMS:
                by_model.setdefault(e.get("model_id"), []).append((c, b))
    for group in by_model.values():
        if len(group) < 2:
            continue
        keeper = max(range(len(group)), key=lambda i: (_chart_matches_section(*group[i]),
                                                        len((ci.get(group[i][0]) or {}).get("refs") or ())))
        for i, (c, b) in enumerate(group):
            if i != keeper and c in b.chart_ids:
                b.chart_ids.remove(c)
                b.origin.pop(c, None)


def _greedy_diverse(chosen: list[str], pool: list[str], ci: dict, target: int) -> list[str]:
    """Take from `pool` until `chosen` reaches `target`, each step preferring the LEAST-used model,
    then the least-used role, then the earliest candidate.

    Counts, not membership. A set-membership test ("is this model already present?") stops
    discriminating the moment every model appears once, and then falls back to catalog order — which
    clumps, because a model's charts are adjacent in the index. That reproduced the exact defect:
    6 charts of which 4 were one model. Counting keeps balancing all the way to the ceiling, which is
    what "don't show four views of one decomposition" (6/8 reviews) actually requires.
    """
    order = list(ci)
    pool = list(pool)
    while len(chosen) < target and pool:
        rc = collections.Counter((ci[c] or {}).get("role") for c in chosen)
        mc = collections.Counter((ci[c] or {}).get("model_id") for c in chosen)
        pool.sort(key=lambda c: (mc[(ci[c] or {}).get("model_id")],
                                 rc[(ci[c] or {}).get("role")],
                                 order.index(c)))
        chosen.append(pool.pop(0))
    return chosen


def _exhibit_contract(picked: list[str], ci: dict) -> list[str]:
    """Return between _CHART_FLOOR and _CHART_CEILING chart ids. Deterministic.

    Three cases, all diversity-aware:
      within band  -> take the planner's picks as they are; that judgement is worth keeping
      over ceiling -> choose from ITS picks, maximising role then model diversity. Blind truncation
                      would take the first N, which are typically all one model — reproducing the
                      `redundant-exhibits` defect (6/8 reviews; the treasurer piece ran four charts of
                      one decomposition) while nominally obeying the ceiling.
      under floor  -> keep every pick, fill from the rest the same way.
    Near-duplicates (same model, same refs) are dropped throughout.
    """
    chosen: list[str] = []
    seen: set = set()
    per_model: collections.Counter = collections.Counter()
    for cid in [c for c in picked if c in ci]:            # planner's picks: dedupe + cap per model
        k, mid = _chart_key(cid, ci), (ci[cid] or {}).get("model_id")
        if k in seen or per_model[mid] >= _MAX_PER_MODEL:
            continue
        seen.add(k)
        per_model[mid] += 1
        chosen.append(cid)

    if len(chosen) > _CHART_CEILING:
        return _greedy_diverse([], chosen, ci, _CHART_CEILING)

    def _pool(cap: int | None) -> list[str]:
        out, s, pm = [], set(seen), collections.Counter(per_model)
        for c in ci:
            k, mid = _chart_key(c, ci), (ci[c] or {}).get("model_id")
            if k in s or (cap is not None and pm[mid] >= cap):
                continue
            s.add(k)
            pm[mid] += 1
            out.append(c)
        return out

    chosen = _greedy_diverse(chosen, _pool(_MAX_PER_MODEL), ci, _CHART_FLOOR)
    if len(chosen) < _CHART_FLOOR:                        # relax the cap rather than ship under floor
        already = {_chart_key(c, ci) for c in chosen}
        chosen = _greedy_diverse(chosen, [c for c in ci if _chart_key(c, ci) not in already],
                                 ci, _CHART_FLOOR)
    return chosen


def _resolve_chart_id(cid: str, chart_index: dict) -> str | None:
    """Map a planner's chart reference to a real chart id — exact, else substring, else closest match.
    The planner sometimes paraphrases the id ('the funding stack' vs the exact title); recover it rather
    than silently drop the chart (which left one persona chartless)."""
    if cid in chart_index:
        return cid
    low = (cid or "").lower().strip()
    if low:
        for k in chart_index:
            kl = k.lower()
            if low == kl or (len(low) > 6 and (low in kl or kl in low)):
                return k
    m = difflib.get_close_matches(cid, list(chart_index), n=1, cutoff=0.55)
    return m[0] if m else None


_STUDIO_CACHE: dict[tuple, str] = {}


def _studio_chart(brief: dict, cid: str, info: dict, run: dict, out_dir: Path, prose: str = "") -> str | None:
    """Design ONE article chart through the agentic Chart Studio — the SoTA multimodal critique loop
    (framer → proposer → vision critique → reviser → judge), PROSE-DRIVEN so the form matches what the
    section says the chart shows. Returns a base64 PNG, or None to fall back to the deterministic family
    renderer for that one chart. This is the default build path; disable the whole studio with
    ARTICLE_CHART_STUDIO=0. Process-cached by (model, chart, as_of) to bound the vision-loop cost."""
    key = (info.get("model_id"), cid, str(getattr(run.get("latest"), "as_of", "")))
    if key in _STUDIO_CACHE:
        return _STUDIO_CACHE[key]
    try:
        from .studio.from_model import brief_for_chart
        from .studio.graph import run_studio
        p = brief["mat"]["p"]
        ib = brief_for_chart(p["name"], p.get("decision", ""), info["model_id"], cid, run, prose=prose)
        if ib is None:
            return None
        st = run_studio(ib, str(out_dir), max_iterations=2)   # 2 vision cycles — catch label/whitespace defects
        png = st.get("png_path")
        if png and Path(png).exists():
            b64 = base64.b64encode(Path(png).read_bytes()).decode()
            _STUDIO_CACHE[key] = b64
            return b64
    except Exception as exc:
        print(f"STUDIO fallback — {cid}: {type(exc).__name__}: {str(exc)[:60]}", file=sys.stderr)
    return None


def _render_charts(brief: dict, chart_ids: list[str], out_dir: Path,
                   prose_by_cid: dict | None = None) -> dict[str, tuple[Path, str]]:
    """Render each referenced chart once → {chart_id: (png_path, caption)}. The Chart Studio is the
    default build path (prose-driven form + vision critic); the deterministic family renderer is the
    per-chart fallback when the Studio errs or is disabled (ARTICLE_CHART_STUDIO=0)."""
    runs = brief["mat"]["runs"]
    studio_on = os.environ.get("ARTICLE_CHART_STUDIO", "1") != "0"
    prose_by_cid = prose_by_cid or {}
    out: dict[str, tuple[Path, str]] = {}
    for i, cid in enumerate(dict.fromkeys(chart_ids)):
        info = brief["chart_index"].get(cid)
        if not info:
            continue
        run = runs.get(info["model_id"]) or {}
        b64, cap = None, None
        if studio_on:
            b64 = _studio_chart(brief, cid, info, run, out_dir, prose=prose_by_cid.get(cid, ""))
            cap = decisive(info.get("insight", "") or cid) if b64 else None
        if not b64:
            b64, cap = chart_png_family(run, cid)
        if not b64:
            b64 = chart_png(run, cid)
            cap = decisive(info.get("insight", "") or cid)
        if b64:
            pth = out_dir / f"chart_{i}.png"
            pth.write_bytes(base64.b64decode(b64))
            out[cid] = (pth, _docx_caption(cap or info.get("insight", "")))
    return out


def _inject_cross_jurisdiction(persona_id: str, mat: dict, conn, out_dir: Path,
                               filled_sections: list, charts: dict) -> str | None:
    """Render ONE cross-country exhibit for the persona's first jurisdiction-generic model and place it in
    the section that discusses it (else the last section). Defensive — any failure just means no bonus
    chart. Uses the deterministic multi-line render, and the freshness guard drops stale jurisdictions."""
    if conn is None:
        return None
    from .studio.from_model import GRAPH_DIR as _GD
    from .studio.from_model import studio_cross_jurisdiction
    models = (mat.get("p") or {}).get("models") or []
    for mid in models:
        try:
            d = yaml.safe_load((_GD / f"{mid}.yaml").read_text())
        except Exception:
            continue
        if not (d.get("instances") and (d.get("generic_over") or len(d.get("instances")) > 1)):
            continue
        chart = next((c for c in d.get("charts", [])
                      if (c.get("data_contract", {}) or {}).get("kind") not in ("scatter", "pearson")), None)
        if not chart:
            continue
        xdir = out_dir / "_xjur"
        r = studio_cross_jurisdiction(mid, chart["id"], conn, str(xdir),
                                      persona=(mat["p"].get("name") or persona_id),
                                      decision=(mat["p"].get("decision") or ""))
        png = r.get("png")
        if not png or not Path(png).exists():
            continue
        dest = out_dir / "chart_xjur.png"
        dest.write_bytes(Path(png).read_bytes())
        cid = f"xjur::{mid}"
        charts[cid] = (dest, _docx_caption(
            f"The same model, run across central banks — {chart['id'].lower()}. The divergence between "
            f"economies, not any one line, is the point."))
        # place in the section whose prose names the model's topic, else the last content section
        low_mid = mid.replace("_", " ")
        home = next((s for s in filled_sections if low_mid in (s.prose or "").lower()), None)
        home = home or (filled_sections[-1] if filled_sections else None)
        if home is not None:
            home.chart_ids.append(cid)
            print(f"XJUR — {persona_id}: added cross-country {mid} exhibit to «{home.heading}»", file=sys.stderr)
            return cid
    return None


def _docx_caption(text: str, width: int = 200) -> str:
    """A figure caption for the .docx: whole words only, ellipsis at a word boundary — never the bare
    `[:160]` mid-word chop that shipped 'truncated captions' the agency flagged as publish seams."""
    t = " ".join((text or "").split())
    return textwrap.shorten(t, width=width, placeholder="…") if t else t


# ── semantic chart↔section binding (A) ──────────────────────────────────────────────────────────
def _match_plan(sec: SectionDraft, plans: list, used: set):
    """The planned section a drafted section came from — by heading similarity, tie-broken by the
    plan's model being named in the prose. Greedy over the unused plans, so it survives the writer
    renaming, reordering, or adding sections across the resampling loop (index alignment does not)."""
    cands = [pl for pl in plans if id(pl) not in used]
    if not cands:
        return None

    def score(pl) -> float:
        h = difflib.SequenceMatcher(None, (sec.heading or "").lower(), (pl.heading or "").lower()).ratio()
        mid = (pl.model_id or "").replace("_", " ").lower()
        return h + (0.15 if mid and mid in (sec.prose or "").lower() else 0.0)

    return max(cands, key=score)


def _bind_sections(outline, filled_sections: list[SectionDraft], ci: dict) -> list[SectionBinding]:
    """Pair each drafted section with the planner's chart picks and model_id for that section. Returns
    bindings 1:1 with `filled_sections` (same order), so placement writes back cleanly."""
    plans, used, bindings = list(outline.sections), set(), []
    for sec in filled_sections:
        plan = _match_plan(sec, plans, used)
        ids: list[str] = []
        mid = ""
        if plan is not None:
            used.add(id(plan))
            mid = plan.model_id or ""
            for c in plan.chart_ids:
                r = _resolve_chart_id(c, ci)
                if r and r not in ids:
                    ids.append(r)
        bindings.append(SectionBinding(heading=sec.heading, model_id=mid, prose=sec.prose,
                                       chart_ids=ids, origin={c: "planned" for c in ids}))
    return bindings


def _dedup_keep(cids: list[str], ci: dict) -> list[str]:
    """Order-preserving drop of near-duplicate charts (same model + refs), no truncation."""
    out, seen = [], set()
    for c in cids:
        if c not in ci:
            continue
        k = _chart_key(c, ci)
        if k not in seen:
            seen.add(k)
            out.append(c)
    return out


def _place_charts(bindings: list[SectionBinding], ci: dict, *, soft_cap: int = 16) -> None:
    """Enforce the exhibit contract on the charts the sections ASKED FOR, then return each surviving
    chart to its home section — by identity, never by `i % len`.

    Owner choice (2026-07-17): the argument drives the count. Every chart a section named is kept, the
    inline ceiling flexing up to `soft_cap`; only past the cap do we diversify-truncate. Under the
    floor, `_exhibit_contract` fills from the pool as before. A floor-filled chart lands in the section
    whose model owns it, because every chart knows its model and every section declares one.
    """
    if not ci or not bindings:
        return
    picked: list[str] = []
    for b in bindings:
        for c in b.chart_ids:
            if c not in picked:
                picked.append(c)

    if len(picked) < _CHART_FLOOR:
        final = _exhibit_contract(picked, ci)                      # under floor → fill up to the floor
    elif len(picked) <= soft_cap:
        final = _dedup_keep(picked, ci)                            # 4..cap → keep what the argument named
    else:
        final = _greedy_diverse([], picked, ci, soft_cap)          # over the cap → diversify down to it
    if len(final) != len(picked):
        print(f"EXHIBIT CONTRACT — sections named {len(picked)}, shipping {len(final)} "
              f"(floor {_CHART_FLOOR}, soft cap {soft_cap})", file=sys.stderr)

    finalset = set(final)
    for b in bindings:                                             # 1) keep each section's own survivors
        b.chart_ids = [c for c in b.chart_ids if c in finalset]
        b.origin = {c: b.origin.get(c, "planned") for c in b.chart_ids}
    placed = {c for b in bindings for c in b.chart_ids}
    for cid in final:                                              # 2) home the floor-fill by provenance
        if cid in placed:
            continue
        mid = (ci.get(cid) or {}).get("model_id")
        home = next((b for b in bindings if b.model_id and b.model_id == mid), None) \
            or min(bindings, key=lambda b: len(b.chart_ids))
        home.chart_ids.append(cid)
        home.origin[cid] = "floor"
        placed.add(cid)

    # dedup AFTER placement: within a section collapse a subset chain (the double-EBP, a locus + its
    # two-series twin) keeping the prose form; across sections drop only near-exact duplicates. A
    # cross-section subset (a gap chart and its parent fan) is never dropped — the v4 Fed regression.
    dropped = _dedup_bindings(bindings, ci)
    if dropped:
        print(f"CONCEPT-DEDUP charts — dropped {len(dropped)} redundant: {', '.join(dropped)[:90]}",
              file=sys.stderr)
    _spread_overloaded_sections(bindings, ci)


def _chart_matches_section(cid: str, sec: SectionBinding) -> int:
    words = set(_distinctive_words(cid))
    pl = (sec.prose or "").lower()
    return sum(1 for w in words if w in pl)


def _spread_overloaded_sections(bindings: list[SectionBinding], ci: dict, *, cap: int = 4) -> None:
    """No section should carry a third of the article's charts — the v5 openings stuffed six. Move the
    weakest-matching excess charts off an over-cap section to the best-matching under-cap section, so
    charts land where their prose reads them, not clumped in the opener."""
    for b in bindings:
        guard = 0
        while len(b.chart_ids) > cap and guard < 20:
            guard += 1
            others = [o for o in bindings if o is not b and len(o.chart_ids) < cap]
            if not others:
                break

            def _gain(cid):                                   # how much better this chart fits elsewhere
                return max(_chart_matches_section(cid, o) for o in others) - _chart_matches_section(cid, b)

            mover = max(b.chart_ids, key=_gain)               # the chart most out of place here
            dest = max(others, key=lambda o: (_chart_matches_section(mover, o), -len(o.chart_ids)))
            b.chart_ids.remove(mover)
            dest.chart_ids.append(mover)
            dest.origin[mover] = b.origin.pop(mover, "planned")


# ── prose↔chart completeness (B) ────────────────────────────────────────────────────────────────
# A NAMED EXHIBIT is a specific visual the reader is told to expect — "the Phillips curve", "the
# Beveridge curve", "the Sahm rule", "the variance risk premium". Its telltale is a proper phrase
# ending in an exhibit noun. This must be high-PRECISION: an article is grounded in its models, so it
# legitimately mentions their topics (inflation, the output gap, the reaction function) — matching on
# topic words pulls in every chart (the central banker "named" 19). We fire only on named exhibits, so
# a chart is built when the prose promises a picture, not merely when it discusses a subject.
_EXHIBIT_BASE = {"curve", "index", "rule", "premium", "surface", "ladder", "quadrant", "fan",
                 "arch", "distribution", "cone", "band", "heatmap", "frontier", "dumbbell"}
# v6: charts the prose named as a FORM but the gate never recognised, so they went "described but not
# drawn" for rounds — the recession probit (Fed), the realized-vs-implied scatter (vol). These extend
# the phrase-matching set BUT are kept OUT of the distinctive-word strip below (unlike the generic base,
# "scatter"/"probit" are specific enough to stay distinctive for _prose_describes).
_EXHIBIT_NOUN = _EXHIBIT_BASE | {"probit", "scatter"}


def _named_phrases(cid: str, info: dict) -> list[str]:
    """The exhibit phrases that would identify this chart in prose: the model spelled out, plus any
    trailing 'X curve / Y index' phrase in the chart's own title — each ending in an exhibit noun."""
    out = []
    mid = (info.get("model_id") or "").replace("_", " ").strip().lower()
    if mid and any(w in _EXHIBIT_NOUN for w in mid.split()):
        out.append(mid)
    title = (cid or "").lower()
    for m in re.finditer(r"([a-z]+(?:[ -][a-z]+)?[ -](?:%s))" % "|".join(_EXHIBIT_NOUN), title):
        out.append(m.group(1).replace("-", " ").strip())
    return out


def _prose_names(cid: str, info: dict, low: str) -> bool:
    """Does the finished prose name this chart as an exhibit? True only when one of its exhibit phrases
    appears verbatim — high precision, so completeness never stuffs the article with topic charts."""
    return any(len(ph) >= 6 and ph in low for ph in _named_phrases(cid, info))


# v4/N2: the prose also DESCRIBES charts that aren't proper-noun exhibits — "the slope chart", "the
# realized-vs-implied scatter", "the now-vs-quarter curve". Catch them without the v3 over-match by
# requiring a sentence that BOTH points at a visual AND names ≥2 of the chart's distinctive words.
_VISUAL = re.compile(r"\b(chart|charts|scatter|plot|plots|panel|exhibit|figure|graph|shows?|showing|"
                     r"decomposition|heatmap|surface)\b", re.I)
_CHART_STOP2 = {"over", "time", "its", "versus", "through", "against", "today", "each", "both", "into",
                "from", "with", "this", "that", "across", "between", "than", "and", "the", "now"}


def _distinctive_words(cid: str) -> list[str]:
    return [w for w in re.findall(r"[a-z][a-z-]{3,}", (cid or "").lower())
            if w not in _CHART_STOP2 and w not in _EXHIBIT_BASE]


def _prose_describes(cid: str, sents: list[str]) -> bool:
    """A described (non-exhibit) chart: some sentence points at a visual AND names ≥2 distinctive words
    of the chart. Conservative — two words + a visual verb, so a bare topic mention never fires."""
    words = _distinctive_words(cid)
    if len(words) < 2:
        return False
    return any(_VISUAL.search(s) and sum(1 for w in words if w in s) >= 2 for s in sents)


def _section_for_chart(bindings: list[SectionBinding], info: dict, low: str,
                       cid: str = "") -> SectionBinding | None:
    """The section a chart belongs to: grounded in its model, else the section whose PROSE best names
    the chart's subject, else the LEAST-LOADED section. NEVER the opening by default — that front-loaded
    four v5 articles (a chart that couldn't be homed fell to bindings[0], stuffing the first section)."""
    mid = info.get("model_id") or ""
    home = next((b for b in bindings if b.model_id and b.model_id == mid), None)
    if home is not None:
        return home
    if mid:
        home = next((b for b in bindings if mid.replace("_", " ") in (b.prose or "").lower()), None)
        if home is not None:
            return home
    # score sections by how much of the chart's subject their prose actually discusses. The index is the
    # final tiebreaker so two equally-scored, equally-loaded sections never fall through to comparing the
    # (unorderable) SectionBinding objects themselves.
    words = set(_distinctive_words(cid)) | {w for w in re.findall(r"[a-z]{4,}", (info.get("insight") or "").lower())}
    best = max(range(len(bindings)),
               key=lambda i: (sum(1 for w in words if w in (bindings[i].prose or "").lower()),
                              -len(bindings[i].chart_ids), -i),
               default=None)
    if best is not None and sum(1 for w in words if w in (bindings[best].prose or "").lower()) > 0:
        return bindings[best]
    return min(bindings, key=lambda b: len(b.chart_ids)) if bindings else None


def _reconcile_prose_charts(full_text: str, bindings: list[SectionBinding], brief: dict,
                            conn=None, mat=None, *, soft_cap: int = 16) -> list[str]:
    """Build every chart the finished prose names. The writer names charts in words (the Phillips
    curve, the Beveridge curve); nothing used to reconcile those words with the built set, so a chart
    the argument builds toward could simply not exist. Scan the prose against the persona's chart
    universe, and attach each named-but-unbuilt chart to the section that names it — pulling and
    executing its model if selection had left it out. Returns the chart ids added."""
    ci = brief["chart_index"]
    low = (full_text or "").lower()
    sents = re.split(r"(?<=[.!?])\s+", low)
    built = {c for b in bindings for c in b.chart_ids}

    # 1) prose-named (exhibit) OR prose-described (visual + ≥2 distinctive words) charts, in the index
    named = [cid for cid, info in ci.items()
             if _prose_names(cid, info, low) or _prose_describes(cid, sents)]

    # 2) prose-named charts whose model selection dropped entirely — pull and run it (best effort)
    if conn is not None and mat is not None:
        ran = set(mat.get("runs", {}))
        try:
            candidates = yaml.safe_load((GRAPH_DIR / "personas.yaml").read_text())["personas"][mat["id"]]["models"]
        except Exception:
            candidates = []
        for mid in candidates:
            if mid in ran or mid.replace("_", " ") not in low:
                continue
            try:
                run = graph_corpus.run_model(mid, conn)
            except Exception as exc:
                print(f"PULL FAILED — {mid}: {type(exc).__name__}: {str(exc)[:60]}", file=sys.stderr)
                continue
            mat["runs"][mid] = run
            for c in (run.get("charts") or []):
                cid = c.get("id")
                if not cid:
                    continue
                ci[cid] = {"model_id": mid, "insight": c.get("insight", ""),
                           "refs": c.get("data_contract", {}) or {}}
                if _prose_names(cid, ci[cid], low) or _prose_describes(cid, sents):
                    named.append(cid)
            print(f"PULLED — {mat['id']}: {mid} (prose names it; selection had dropped it)", file=sys.stderr)

    added = []
    for cid in named:
        if cid in built:
            continue
        home = _section_for_chart(bindings, ci[cid], low, cid)
        if home is not None:
            home.chart_ids.append(cid)
            home.origin[cid] = "prose"
            built.add(cid)
            added.append(cid)
    if len(built) > soft_cap:
        print(f"OVER-CAP — {mat['id'] if mat else '?'}: prose names {len(built)} charts (soft cap "
              f"{soft_cap}); placement will diversify-truncate and a human must reconcile the prose.",
              file=sys.stderr)
    return added


def _assemble_docx(path: Path, p: dict, mat: dict, headline: str, standfirst: str, exec_summary: str,
                   sections: list[SectionDraft], charts: dict, ill_png: Path, infog_png: Path | None) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    def _fig(png: Path, caption: str = "", width: float = 6.5) -> None:
        doc.add_picture(str(png), width=Inches(width))
        img_para = doc.paragraphs[-1]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            # bind the image to its caption so a page break can't strand the caption on the next page
            # (the 'orphaned captions' the agency flagged) — keep_with_next holds the pair together.
            img_para.paragraph_format.keep_with_next = True
            c = doc.add_paragraph(caption)
            try:
                c.style = doc.styles["Caption"]
            except KeyError:
                pass
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.keep_together = True

    doc = Document()
    doc.add_heading(headline, level=0)
    sf = doc.add_paragraph()
    r = sf.add_run(standfirst); r.italic = True; r.font.size = Pt(13)
    by = _byline()
    dl = doc.add_paragraph()
    if by.get("author"):                                   # human byline above the fold
        who = by["author"] + (f", {by['credential']}" if by.get("credential") else "")
        dl.add_run(f"By {who}  ·  ").bold = True
    dl.add_run(f"{p['name']}.  Data as of {mat.get('as_of','')}.").italic = True
    if by.get("tagline"):                                  # the one-line credential under the byline
        tl = doc.add_paragraph()
        tr = tl.add_run(by["tagline"]); tr.italic = True; tr.font.size = Pt(9.5)

    _fig(ill_png)                                          # the Van Gogh header

    # the DUAL executive summary: narrative prose, then the infographic as the visual summary
    for para in [x for x in exec_summary.split("\n") if x.strip()]:
        doc.add_paragraph(para.strip())
    if infog_png and Path(infog_png).exists():
        _fig(infog_png, "The whole picture at a glance — every reading, priced and traced.")

    for s in sections:
        doc.add_heading(s.heading, level=2)
        for para in [x for x in s.prose.split("\n") if x.strip()]:
            doc.add_paragraph(para.strip())
        for cid in s.chart_ids:
            if cid in charts:
                cp, cap = charts[cid]
                _fig(cp, cap, width=6.0)

    resolved = _resolve_papers(mat.get("papers", []))
    if resolved:
        doc.add_heading("Further reading", level=2)
        for r in resolved:
            para = doc.add_paragraph(style="List Bullet")
            if r["url"]:
                _add_hyperlink(para, r["url"], r["title"])
            else:
                para.add_run(r["title"]).bold = True
            tail = " — " + (f"{r['byline']}. " if r["byline"] else "") + (r["summary"] or "")
            para.add_run(tail)

    # the pipeline hook: a soft About-the-author + call-to-action, after the analysis so it never
    # interrupts the read (the lead-gen gap the marketing review named). Static prose — no figure tokens.
    if by.get("bio") or by.get("cta_text"):
        doc.add_heading("About the author", level=2)
        if by.get("bio"):
            bp = doc.add_paragraph()
            who = by.get("author", "")
            # only lead with the name when the bio doesn't already open with it (third-person bios do)
            if who and not by["bio"].lstrip().lower().startswith(who.split()[0].lower()):
                bp.add_run(f"{who}. ").bold = True
            bp.add_run(by["bio"])
        if by.get("cta_text"):
            cta = doc.add_paragraph()
            r = cta.add_run(by["cta_text"] + " "); r.italic = True
            if by.get("cta_url"):
                _add_hyperlink(cta, by["cta_url"], by.get("cta_link_text") or by["cta_url"])
            elif by.get("cta_link_text"):
                cta.add_run(by["cta_link_text"]).italic = True

    foot = doc.add_paragraph()
    fr = foot.add_run(f"Source: {', '.join(mat.get('source_labels', [])) or 'UMD'}. "
                      f"Models: {', '.join(mat.get('model_names', []))}. Every figure is executed on data; "
                      f"the prose is authored, the numbers are not.")
    fr.font.size = Pt(8); fr.italic = True
    doc.save(str(path))


# ── the pipeline stages (node bodies for the article graph) ──────────────────────────────────────
def _rank(leak_free: bool, is_grounded: bool, is_crit_ok: bool) -> int:
    """GROUNDING OUTRANKS STYLE. Said plainly everywhere and coded backwards the first time: flat
    scores let a critic-clean-but-unproven draft (4) beat a grounded-but-clumsy one (3), and
    volatility_trader duly shipped "rank 4/4: lints clean, critic ok" with grounded=False. A graceful
    sentence the arithmetic cannot confirm is worth less than an awkward one it can."""
    if not leak_free:
        return 0                                       # an untraced figure is the worst outcome here
    return 1 + (2 if is_grounded else 0) + (1 if is_crit_ok else 0)


def draft_best(persona_id: str, brief: dict, outline, conn, *, max_iter: int = 3) -> dict:
    """The best-of-N redraft loop: write → firewall → lint gates → ground (Judge) → style critic,
    keeping the highest-ranked draft across iterations (grounding outranks style). One cohesive node:
    the draft / ground / critique steps form an atomic resample loop and must stay together."""
    feedback, reasons, crit_ok = "", [], False
    grounded, judge_failures = False, []
    filled_sections, headline = [], outline.headline
    standfirst, exec_summary, full_text = outline.standfirst, "", ""
    # KEEP THE BEST DRAFT, NOT THE LAST ONE. Each iteration is a full rewrite, so a draft that fixes
    # the leak can introduce slop and vice versa — it resamples rather than converges. The loop simply
    # fell out at max_iter and shipped whatever the final pass produced: on 2026-07-17 that was
    # ['iter0: leak', 'iter1: ungrounded', 'iter2: leak'], so an UNTRACED FIGURE shipped while a
    # cleaner draft from iter1 was thrown away. Rule #2 lost to a loop-control detail.
    best: tuple[int, dict] | None = None

    def _keep(score: int, **snap) -> None:
        nonlocal best
        if best is None or score > best[0]:
            best = (score, snap)

    for it in range(max_iter):
        article = write_article(brief, outline, feedback=feedback)
        full_text, standfirst, exec_summary, filled_sections, leak = _finalize_article(article, brief["toks"])
        slop = _slop_lint(full_text)
        stray = re.search(r"\{[^}]*\}", full_text)          # an invented figure marker / bad token
        epi = _episode_leaks(full_text, brief.get("data_start", ""))
        wordnum = _worded_precision(full_text)
        if leak or slop or stray or epi or wordnum:
            bits = []
            if leak:
                # Two ways out, and the writer needs to be told BOTH — otherwise it rewrites the
                # sentence, needs the number again, types it again, and the loop burns its whole
                # budget. On 2026-07-17 all three CB drafts leaked and the Judge never ran: the
                # figures were "more than 100bp tighter" (true — the gap is ~140bp), "the 2% target"
                # (a model parameter) and "around 1%" (r* is 1.06). Every one a true statement the
                # gate had no legitimate way to express.
                bits.append(
                    f"The figure '{leak}' is typed in digits but is not a token. Digits are reserved "
                    f"for measured values. EITHER cite it as its {{n}} token if it is a real executed "
                    f"figure, OR — if you meant it loosely (a rough magnitude, a round threshold, a "
                    f"target) — SPELL IT OUT IN WORDS instead: 'more than a hundred basis points', "
                    f"'the two percent target', 'around one percent'. Do not simply delete the point. "
                    f"But choose ONE: the words REPLACE the figure — never put the spelled-out number "
                    f"next to its own token or numeral (never 'above five percent {{n}}', never 'around "
                    f"one percent 1.06%'), and never cite the same token twice in a row.")
            if wordnum:
                bits.append(f"REWRITE '{wordnum}': you have spelled out a precise value in words to "
                            f"avoid a token. Words are for approximations only. If the figure is real, "
                            f"cite its {{n}} token; if you meant roughly, round it and say so.")
            if stray:
                bits.append("REMOVE every {…} marker from the prose — do not number or footnote the "
                            "charts; name a chart in words. The only braces allowed are the number tokens.")
            if slop:
                bits.append("DELETE these AI-slop phrases and rewrite the sentence plainly: "
                            + ", ".join(f"'{s}'" for s in slop))
            if epi:
                sy = brief.get("data_start", "")[:4]
                bits.append(f"REMOVE every reference to {', '.join(epi)}: the data begins {sy}, so no chart "
                            f"shows those years — naming them fabricates history. Rewrite using ONLY episodes "
                            f"on or after {sy} (e.g. what the series actually spans).")
            feedback = " ".join(bits)
            reasons.append(f"iter{it}: {f'leak({leak}) ' if leak else ''}{'stray ' if stray else ''}"
                           f"{'slop ' if slop else ''}{'episode ' if epi else ''}"
                           f"{f'worded({wordnum})' if wordnum else ''}".strip())
            # A draft with an untraced figure is the worst thing here — it puts a number in front of a
            # reader that no model produced. Rank it below everything, but still keep it: if every
            # draft leaks, the best of a bad set must ship with the failure reported, not vanish.
            _keep(_rank(not leak, False, False), full_text=full_text, standfirst=standfirst,
                  exec_summary=exec_summary, filled_sections=filled_sections,
                  grounded=False, judge_failures=[], crit_ok=False)
            continue
        # §06 role 7. Grounding is adjudicated BEFORE style: a graceful sentence that contradicts the
        # model is worse than a clumsy one that does not, and the style critic cannot tell the
        # difference — it has never seen a number. The LLM extracts the claims; arithmetic settles them
        # against the very rows this brief was written from.
        grounded, judge_failures = _judge(full_text, brief, conn)
        failures = judge_failures
        # Lints clean. Better than any leaking draft, and better still if the arithmetic agrees.
        _keep(_rank(True, grounded, False), full_text=full_text, standfirst=standfirst,
              exec_summary=exec_summary, filled_sections=filled_sections,
              grounded=grounded, judge_failures=judge_failures, crit_ok=False)
        if not grounded and failures:
            feedback = ("These sentences CONTRADICT the executed model output. Each is followed by the "
                        "arithmetic that settles it, taken from the model's own series. Rewrite each one "
                        "to say what the numbers say — do not soften it, do not delete the paragraph, and "
                        "do not invent a replacement figure: "
                        + " | ".join(f"{v.quote!r} — {v.detail}" for v in failures[:5]))
            reasons.append(f"iter{it}: ungrounded {len(failures)} claims")
            continue
        crit = critique_article(full_text, headline)
        if crit.ok:
            crit_ok = True
            _keep(_rank(True, grounded, True), full_text=full_text, standfirst=standfirst, exec_summary=exec_summary,
                  filled_sections=filled_sections, grounded=grounded,
                  judge_failures=judge_failures, crit_ok=True)
            break
        feedback = "; ".join(crit.fixes or crit.defects)
        reasons.append(f"iter{it}: critic {len(crit.defects)} defects")

    # Whatever the loop's last pass happened to produce, ship the best draft it ever saw.
    if best is not None:
        snap = best[1]
        full_text, standfirst = snap["full_text"], snap["standfirst"]
        exec_summary, filled_sections = snap["exec_summary"], snap["filled_sections"]
        grounded, judge_failures, crit_ok = snap["grounded"], snap["judge_failures"], snap["crit_ok"]
        reasons.append(f"shipped the best of {max_iter} drafts (rank {best[0]}/4: "
                       f"{'lints clean' if best[0] >= 1 else 'ALL DRAFTS LEAK'}"
                       f"{', grounded' if grounded else ', NOT GROUNDED'}"
                       f"{', critic ok' if crit_ok else ''})")
    if not grounded:
        print(f"NOT GROUNDED — {persona_id}: shipping prose the arithmetic does not confirm. "
              f"{len(judge_failures)} claim(s) contradict the executed models. A human must read this.",
              file=sys.stderr)

    # final safety: strip any residual marker so nothing broken ships
    standfirst = _strip_stray(standfirst)
    exec_summary = _strip_stray(exec_summary)
    for s in filled_sections:
        s.prose = _strip_stray(s.prose)
    return {"full_text": full_text, "standfirst": standfirst, "exec_summary": exec_summary,
            "filled_sections": filled_sections, "grounded": grounded, "judge_failures": judge_failures,
            "crit_ok": crit_ok, "headline": headline, "reasons": reasons}


def reconcile_charts(brief: dict, outline, draft: dict, conn, mat: dict) -> dict:
    """A + B: bind each section to the charts it reads, build every chart the prose names, place them
    by identity. Mutates the draft's sections' chart_ids; returns the bindings and any added reasons."""
    ci = brief["chart_index"]
    filled_sections = draft["filled_sections"]
    bindings = _bind_sections(outline, filled_sections, ci)
    added = _reconcile_prose_charts(draft["full_text"], bindings, brief, conn=conn, mat=mat)
    reasons = list(draft.get("reasons", []))
    if added:
        reasons.append(f"built {len(added)} prose-named chart(s): {', '.join(added)[:80]}")
    _place_charts(bindings, ci)
    for sec, b in zip(filled_sections, bindings):
        sec.chart_ids = list(b.chart_ids)
    return {"bindings": bindings, "reasons": reasons}


def make_illustration(mat: dict, persona_id: str, out_dir: Path, backend: str) -> tuple[Path, dict]:
    """The Van Gogh header — a meaning-carrying metaphor of the article's own finding."""
    finding = _fill_template(mat)
    ill_b64, ill_meta = vangogh.illustration_png(
        finding, title=mat["p"]["title"], decision=mat["p"].get("decision", ""),
        cache_key=f"{persona_id}|article", backend=backend)
    ill_png = out_dir / "illustration.png"
    ill_png.write_bytes(base64.b64decode(ill_b64))
    return ill_png, ill_meta


def reconcile_dashboard(persona_id: str, conn, draft: dict, brief: dict, out_dir: Path) -> dict:
    """C: the dashboard is a projection of THIS finished article — the numbers the prose actually cited
    (so no tile shows a figure the body never used) and the article's own closing read."""
    full_text = draft["full_text"]
    cited_keys = [no.source for no in brief["toks"].values()
                  if no.rendered() and (no.rendered() in full_text or no.rendered().lstrip("+") in full_text)]
    def _sec_charts(s):
        return s.chart_ids if hasattr(s, "chart_ids") else (s.get("chart_ids", []) if isinstance(s, dict) else [])
    body_chart_ids = {c for s in draft.get("filled_sections", []) for c in _sec_charts(s)}
    article_ctx = {"exec_summary": draft["exec_summary"], "full_text": full_text,
                   "thesis": draft["standfirst"], "cited_keys": cited_keys,
                   "body_chart_ids": body_chart_ids}
    infog_png = out_dir / "infographic.png"
    reasons = list(draft.get("reasons", []))
    fam = FAMILY.get(persona_id, decision_brief)
    try:
        fam.render_persona(persona_id, conn, str(infog_png), article=article_ctx)
    except Exception as exc:
        reasons.append(f"infographic→decision_brief: {str(exc).splitlines()[0][:70]}")
        try:
            decision_brief.render_persona(persona_id, conn, str(infog_png), article=article_ctx)
        except Exception:
            infog_png = None
    return {"infog_png": infog_png, "cited_keys": cited_keys, "reasons": reasons}


def assemble(persona_id: str, mat: dict, brief: dict, draft: dict, ill_png, ill_meta: dict,
             infog_png, out_dir: Path, conn=None) -> dict:
    """Render each placed chart once and assemble the .docx; return the run's result dict."""
    p, filled_sections = mat["p"], draft["filled_sections"]
    all_chart_ids = [cid for s in filled_sections for cid in s.chart_ids]
    prose_by_cid = {cid: s.prose for s in filled_sections for cid in s.chart_ids}  # prose-driven form
    charts = _render_charts(brief, all_chart_ids, out_dir, prose_by_cid)
    # global-macro exhibit: if a persona runs a jurisdiction-generic model, add ONE cross-country chart so
    # the article is literally multi-country, not US-shaped (the agency's "global macro is US macro").
    try:
        _inject_cross_jurisdiction(persona_id, mat, conn, out_dir, filled_sections, charts)
    except Exception as exc:
        print(f"XJUR — {persona_id}: skipped ({type(exc).__name__}: {str(exc)[:70]})", file=sys.stderr)
    docx_path = out_dir / "article.docx"
    _assemble_docx(docx_path, p, mat, draft["headline"], draft["standfirst"], draft["exec_summary"],
                   filled_sections, charts, ill_png, infog_png)
    return {"persona": persona_id, "ok": True, "docx_path": str(docx_path), "headline": draft["headline"],
            "standfirst": draft["standfirst"], "sections": len(filled_sections),
            "words": len(draft["full_text"].split()), "critic_ok": draft["crit_ok"],
            "n_charts": len(charts), "caption": ill_meta.get("caption", ""),
            "full_text": draft["full_text"], "reasons": draft.get("reasons", []),
            # Reported SEPARATELY from critic_ok, and it is the one that matters. critic_ok is a style
            # score; `grounded` is whether the prose agrees with the arithmetic. The shipped CB article
            # was critic_ok=True with two sentences that contradicted their own series.
            "grounded": draft["grounded"], "ungrounded": [v.model_dump() for v in draft["judge_failures"]]}


# ── the orchestrator ─────────────────────────────────────────────────────────────────────────────
def build_article_full(persona_id: str, conn, out_dir, *, backend: str = "auto", max_iter: int = 3,
                       jurisdiction: str = "US", model_ids: list | None = None) -> dict:
    """Compile and run the article StateGraph, returning its result dict. The graph threads a single
    ArticleState through material → brief → plan → draft → reconcile_charts → reconcile_dashboard →
    assemble. `jurisdiction`/`model_ids` are the steering contract — run a decision-maker's models in a
    chosen currency (defaults reproduce the US persona article)."""
    from .article_graph import run_article
    return run_article(persona_id, conn, out_dir, backend=backend, max_iter=max_iter,
                       jurisdiction=jurisdiction, model_ids=model_ids)["result"]

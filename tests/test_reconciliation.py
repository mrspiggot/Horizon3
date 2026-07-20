"""The reconciliation stage: charts, prose, dashboard and numbers as projections of one state.

These lock in the four fixes that took the v2 batch off a converged ~6.8:
  A  charts bind to the SECTION whose prose reads them, surviving redraft reorder/expansion;
  B  every NAMED EXHIBIT the prose promises is built — and topic-mentions do not stuff the article;
  C  the dashboard's thesis/read/tiles come from the finished article, not the static template;
  D  a concept emitted by two models collapses to ONE citable number everywhere.
"""
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO))

from render.infographic.schema import NumberObject  # noqa: E402


def _N(key, val, unit, fmt="{:.2f}"):
    return NumberObject(name=key, value=val, unit=unit, fmt=fmt, source=key, as_of="2026-07-17")


# ── Pass 1 — "one figure, once" (number-rendering guard) ──────────────────────────────────────────
def test_one_figure_once_collapses_doubles_but_spares_distinct_numbers():
    from render.writer import _one_figure_once
    toks = {"n0": _N("n0", 0.18, "pp", "{:+.2f}"), "n1": _N("n1", 5.31, "%"), "n2": _N("n2", 1.06, "%")}
    # doubled token, and a worded round approximation run into its own numeral
    assert _one_figure_once("reads +0.18pp +0.18pp, faintly", toks) == "reads +0.18pp, faintly"
    assert _one_figure_once("a rate above five percent 5.31%;", toks) == "a rate above 5.31%;"
    assert _one_figure_once("r* is only about one percent 1.06%.", toks) == "r* is only about 1.06%."
    # must NOT touch two genuinely distinct adjacent figures, nor a lone worded approximation
    assert _one_figure_once("rose from 1.06% to 5.31%", toks) == "rose from 1.06% to 5.31%"
    assert _one_figure_once("an r* of around one percent anchors it", toks) == "an r* of around one percent anchors it"


def test_number_sweep_v4_forms_phrase_and_descriptor():
    from render.writer import _one_figure_once
    toks = {"n1": _N("n1", -0.52, "σ", "{:+.2f}"), "n2": _N("n2", 0.35, "σ", "{:+.2f}")}
    # worded number + noun phrase + matching numeral (10-year deck): keep the phrase, drop the numeral
    assert _one_figure_once("only fifteen percent risk premium 15%", toks) == "only fifteen percent risk premium"
    # bare state-descriptor jammed into a signed/σ token (fin-conditions): insert a connector
    assert _one_figure_once("sits loose -0.52σ", toks) == "sits loose, at -0.52σ"
    assert _one_figure_once("momentum positive +0.35σ", toks) == "momentum positive, at +0.35σ"
    # clean prose with a real connector is untouched
    assert _one_figure_once("the premium is at 15% of the yield", toks) == "the premium is at 15% of the yield"


# ── Pass 1 — one concept once, extended (cross-model numbers + subset charts) ──────────────────────
def test_concept_registry_merges_same_concept_different_field_names():
    from render.infographic.from_persona import _concept_registry
    numbers = {"funding_quality.ig_all_in_pct": _N("funding_quality.ig_all_in_pct", 5.86, "%"),
               "funding_cost.all_in_pct": _N("funding_cost.all_in_pct", 5.05, "%"),
               "funding_quality.hy_all_in_pct": _N("funding_quality.hy_all_in_pct", 8.10, "%")}
    meanings = {"funding_quality.ig_all_in_pct": "all-in IG funding cost",
                "funding_cost.all_in_pct": "all-in fixed IG funding cost",
                "funding_quality.hy_all_in_pct": "all-in HY funding cost"}
    runs = {m.split(".")[0]: {"meta": {"name": m}, "history": [0] * 300} for m in numbers}
    _c, canon = _concept_registry(numbers, meanings, template_keys=set(), runs=runs)
    assert canon["funding_quality.ig_all_in_pct"] == canon["funding_cost.all_in_pct"]  # 5.86≡5.05
    assert canon["funding_quality.hy_all_in_pct"] != canon["funding_quality.ig_all_in_pct"]  # IG≠HY


def test_dedup_within_section_collapses_but_cross_section_subset_survives():
    """The v4 fix: same-section subset collapses (double-EBP), but a cross-section subset (Fed gap chart
    vs its parent fan) is BOTH kept — and a locus survives over its two-series twin (labour Okun)."""
    from render.writer import SectionBinding, _dedup_bindings
    gz = frozenset({"gz_spread", "expected_default", "excess_bond_premium"})
    okun = frozenset({"gdp_growth", "unemployment_change"})
    ci = {
        "GZ decomposition": {"model_id": "credit_ebp", "refs": gz, "kind": "decomposition"},
        "The excess bond premium": {"model_id": "credit_ebp", "refs": frozenset({"excess_bond_premium"}), "kind": "gap_series"},
        "The prescription-minus-actual gap": {"model_id": "reaction", "refs": frozenset({"taylor", "actual"}), "kind": "gap_series"},
        "The prescription fan": {"model_id": "reaction", "refs": frozenset({"taylor", "balanced", "inertial", "actual"}), "kind": "series"},
        "Okun's Law (growth vs change in unemployment)": {"model_id": "okun", "refs": okun, "kind": "scatter"},
        "The two series Okun relates over time": {"model_id": "okun", "refs": okun, "kind": "series"},
    }
    bindings = [
        SectionBinding(heading="splitting the spread", prose="the GZ decomposition and the excess bond premium",
                       model_id="credit_ebp", chart_ids=["GZ decomposition", "The excess bond premium"]),
        SectionBinding(heading="two verdicts", prose="the prescription-minus-actual gap makes the point",
                       model_id="reaction", chart_ids=["The prescription-minus-actual gap"]),
        SectionBinding(heading="the rule's outcome", prose="the prescription fan across every cycle",
                       model_id="reaction", chart_ids=["The prescription fan"]),
        SectionBinding(heading="the corner", prose="recessions pinned to the top-left corner of the scatter",
                       model_id="okun", chart_ids=["Okun's Law (growth vs change in unemployment)",
                                                   "The two series Okun relates over time"]),
    ]
    _dedup_bindings(bindings, ci)
    kept = {c for b in bindings for c in b.chart_ids}
    assert "The excess bond premium" not in kept          # same-section subset → collapsed
    assert "The prescription-minus-actual gap" in kept    # cross-section subset → KEPT (the Fed fix)
    assert "The prescription fan" in kept
    assert "Okun's Law (growth vs change in unemployment)" in kept   # locus kept over its series twin
    assert "The two series Okun relates over time" not in kept       # (prose says corner/scatter)


def test_chart_sanity_drops_total_among_stack_components():
    """P0 — rule #2 at the chart layer: a stacked chart must never include the TOTAL as a band alongside
    its own components (the GZ spread shipped at ~16pp vs a true ~8)."""
    import pandas as pd, numpy as np
    from render.studio.compile import _drop_total_from_stack
    idx = pd.date_range("2007-01-01", periods=12, freq="MS")
    default, ebp = np.linspace(1, 3, 12), np.linspace(2, 5, 12)
    piv = pd.DataFrame({"gz spread": default + ebp, "expected default": default, "ebp": ebp}, index=idx)
    out = _drop_total_from_stack(piv)
    assert "gz spread" not in out.columns                       # the total is dropped
    assert abs(float(out.sum(axis=1).iloc[-1]) - float((default + ebp)[-1])) < 1e-6  # stack = true spread
    # a genuine 2-component stack (no total among them) is untouched
    piv2 = pd.DataFrame({"expected default": default, "ebp": ebp}, index=idx)
    assert list(_drop_total_from_stack(piv2).columns) == ["expected default", "ebp"]


def test_prefer_form_keeps_gap_chart_when_prose_emphasizes_the_gap():
    from render.writer import _prefer_form
    ci = {"The consequence — how far the Fed is from its own rule":
          {"kind": "gap_series", "refs": frozenset({"taylor", "actual"})},
          "The prescription fan":
          {"kind": "series", "refs": frozenset({"taylor", "balanced", "inertial", "actual"})}}
    gap, fan = "The consequence — how far the Fed is from its own rule", "The prescription fan"
    # prose about the gap → keep the gap, not the richer levels fan (the macro_rates regression)
    assert _prefer_form(gap, fan, ci, "the consequence chart — prescription minus actual — that gap is the trade") == gap
    # neutral prose → the richer fan wins (unchanged tie-break)
    assert _prefer_form(gap, fan, ci, "the rules broadly track the funds rate") == fan


def test_cross_jurisdiction_freshness_guard_drops_stale_bindings(capsys):
    """C1 — rule #5 at the jurisdiction boundary: an instance whose latest point is far behind the freshest
    is refused (JP CPI ends 2021 → out of a Phillips comparison), but a fresh one stays (JP unemployment →
    in a Sahm comparison). Never plot stale data as current."""
    from render.graph_corpus import _drop_stale_instances

    class _Pt:
        def __init__(self, a): self.as_of = a

    phillips = {"US": {"latest": _Pt("2026-06-28")}, "EU": {"latest": _Pt("2026-05-28")},
                "GB": {"latest": _Pt("2026-04-28")}, "JP": {"latest": _Pt("2021-06-28")}}
    kept = _drop_stale_instances(phillips, "phillips_curve")
    assert set(kept) == {"US", "EU", "GB"}                  # JP (dead CPI) refused
    assert "refusing to plot stale data" in capsys.readouterr().err
    sahm = {j: {"latest": _Pt("2026-05-28")} for j in ("US", "EU", "GB", "JP")}
    assert set(_drop_stale_instances(sahm, "sahm_rule")) == {"US", "EU", "GB", "JP"}   # all fresh → all kept


def test_lead_hook_off_when_blank_on_when_filled(tmp_path):
    """Agency round B — the byline + closing 'work with me' block appear only when byline.yaml is filled;
    a blank config ships NOTHING half-written, and a filled one adds the byline, About block and CTA link."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import docx
    import render.writer as W
    hdr = tmp_path / "hdr.png"
    fig, ax = plt.subplots(figsize=(2, 1)); ax.plot([0, 1], [0, 1]); fig.savefig(hdr); plt.close(fig)
    mat = {"as_of": "2026-07-20", "papers": [], "source_labels": ["UMD"], "model_names": ["x"]}

    def _render(byline):
        W._byline = lambda: byline
        out = tmp_path / "a.docx"
        W._assemble_docx(out, {"name": "Rates trader"}, mat, "H", "SF", "Exec.", [], {}, hdr, None)
        return [x.text for x in docx.Document(out).paragraphs], docx.Document(out)

    off_texts, off_doc = _render({})                       # blank → hook OFF
    assert not any(t.startswith("By ") for t in off_texts)
    assert "About the author" not in off_texts

    on_texts, on_doc = _render({"author": "Richard Walker", "credential": "Lucidate",
                                "bio": "Builds model-driven research.", "cta_text": "Work with me —",
                                "cta_url": "https://example.com", "cta_link_text": "get in touch"})
    assert any(t.startswith("By Richard Walker") for t in on_texts)
    assert "About the author" in on_texts
    assert any("hyperlink" in r.reltype for r in on_doc.part.rels.values())
    W._byline.cache_clear() if hasattr(W._byline, "cache_clear") else None


def test_docx_caption_truncates_at_a_word_boundary_not_mid_word():
    """Agency round A — the docx figure caption was a bare [:160] slice that cut mid-word ('publish seams').
    It now ends at a word boundary with an ellipsis, and short captions pass through untouched."""
    from render.writer import _docx_caption
    long = ("The GZ credit spread decomposed into compensation for expected default and the excess bond "
            "premium, a fear gauge that spiked hard at Lehman and again at COVID but sits benign today "
            "right across the whole sample window and beyond, a decomposition every credit desk should read")
    out = _docx_caption(long, width=120)
    assert out.endswith("…")
    words = out.rstrip("…").split()
    assert words[-1].isalpha() and len(words[-1]) >= 2      # last kept word is whole, not a mid-word chop
    assert _docx_caption("A short caption.", width=120) == "A short caption."     # short passes through
    assert _docx_caption("") == ""


def test_probit_and_scatter_are_named_exhibits_but_stay_distinctive():
    """P5.4 — the recession probit was 'described but not drawn' for 3 rounds because 'probit' wasn't an
    exhibit noun. It now is (so _prose_names fires), while staying a distinctive word for _prose_describes."""
    from render.writer import _distinctive_words, _prose_describes, _prose_names
    assert _prose_names("The recession probit", {"model_id": "recession_probit"},
                        "the recession probit puts the odds at sixteen percent")
    # scatter stays distinctive (not stripped like the generic base nouns), so _prose_describes still fires
    assert "scatter" in _distinctive_words("The realized-vs-implied volatility scatter")
    assert _prose_describes("The realized-vs-implied volatility scatter",
                            ["the scatter of realized against implied volatility sits in an odd quadrant"])


def test_dashboard_embed_prefers_a_chart_not_already_in_the_body():
    """P4.2 — the dashboard hero should skip a chart already shown standalone in the body (cost-of-money's
    IG-split-twice), but never end up empty: if every candidate is in the body it keeps the first."""
    from render.infographic.families.decomposition_hero import _find_decomposition
    mat = {"p": {"models": ["m"]},
           "runs": {"m": {"charts": [{"id": "A", "data_contract": {"kind": "decomposition"}},
                                     {"id": "B", "data_contract": {"kind": "decomposition"}}]}}}
    assert _find_decomposition(mat)[1]["id"] == "A"                 # first by default
    assert _find_decomposition(mat, {"A"})[1]["id"] == "B"         # A is in the body → pick B
    assert _find_decomposition(mat, {"A", "B"})[1]["id"] == "A"    # all in body → fall back, never empty


def test_cross_section_collapses_the_same_decomposition_twice_but_spares_gap_and_levels():
    """P2a — the SAME model's over-time decomposition shown in two sections (the double-EBP re-grown by a
    Role-2 import) collapses to one, kept where the prose reads it; a gap chart and its levels fan in two
    sections (the Fed regression) are NOT decomposition forms and must both survive."""
    from render.writer import SectionBinding, _dedup_bindings
    # same GZ decomposition imported into the opener AND argued in the credit section → one survives
    O, A = "credit spread decomposition opener", "credit spread decomposition argued"
    ci = {
        O: {"model_id": "credit_ebp", "kind": "decomposition", "refs": frozenset({"gz", "default", "ebp"})},
        A: {"model_id": "credit_ebp", "kind": "decomposition", "refs": frozenset({"gz", "default", "ebp"})},
        "rates gap": {"model_id": "taylor", "kind": "gap_series", "refs": frozenset({"rule", "actual"})},
        "rates fan": {"model_id": "taylor", "kind": "series",
                      "refs": frozenset({"rule", "actual", "balanced", "inertial"})},
    }
    opener = SectionBinding(heading="opening", prose="the setup and the stakes", chart_ids=[O],
                            origin={O: "planned"})
    credit = SectionBinding(heading="credit", prose="the credit spread decomposition into default and premium",
                            chart_ids=[A], origin={A: "planned"})
    rates1 = SectionBinding(heading="the trade", prose="the gap versus the rule", chart_ids=["rates gap"],
                            origin={"rates gap": "planned"})
    rates2 = SectionBinding(heading="the fan", prose="every rule variant as a levels fan", chart_ids=["rates fan"],
                            origin={"rates fan": "planned"})
    dropped = _dedup_bindings([opener, credit, rates1, rates2], ci)
    assert O in dropped and A not in dropped       # kept in the section that argues it, not the opener
    assert credit.chart_ids == [A]
    assert rates1.chart_ids == ["rates gap"] and rates2.chart_ids == ["rates fan"]   # gap/levels both live


def test_placement_never_defaults_to_opening_and_spreads_overflow():
    """P1 — a chart homes to the section whose prose discusses it, never the opening by default; an
    over-stuffed opener spreads its out-of-place charts to their sections (the v5 front-loading)."""
    from render.writer import SectionBinding, _section_for_chart, _spread_overloaded_sections
    bindings = [SectionBinding(heading="opening", prose="the setup and the stakes", chart_ids=[]),
                SectionBinding(heading="the curve", prose="the yield curve slope and its inversion", chart_ids=[])]
    ci = {"The yield-curve slope": {"model_id": "x", "insight": "the slope of the yield curve"}}
    assert _section_for_chart(bindings, ci["The yield-curve slope"], "", "The yield-curve slope").heading == "the curve"
    # tied score AND tied load across sections must not fall through to comparing SectionBindings (crashed
    # the whole batch with `'<' not supported between instances of 'SectionBinding'`)
    tied = [SectionBinding(heading="a", prose="unrelated", chart_ids=[]),
            SectionBinding(heading="b", prose="unrelated", chart_ids=[])]
    got = _section_for_chart(tied, {"model_id": "z", "insight": "nothing matches here"}, "", "some chart id")
    assert got in tied            # returns a section (least-loaded fallback), does not raise
    # an over-cap opener sheds the chart that belongs elsewhere
    op = SectionBinding(heading="opening", prose="the setup",
                        chart_ids=["c1", "c2", "c3", "c4", "The unemployment recession sahm chart"])
    lab = SectionBinding(heading="labour", prose="unemployment and the recession call and the sahm rule", chart_ids=[])
    ci2 = {c: {"model_id": "m", "insight": ""} for c in op.chart_ids}
    _spread_overloaded_sections([op, lab], ci2, cap=4)
    assert len(op.chart_ids) <= 4
    assert "The unemployment recession sahm chart" in lab.chart_ids


def test_prose_describes_catches_visual_plus_two_distinctive_words():
    from render.writer import _prose_describes
    sents = ["the scatter of realized against implied volatility puts today in an unusual quadrant".lower()]
    assert _prose_describes("The realized-vs-implied volatility scatter", sents) is True
    # a bare topic mention (no visual verb, or <2 distinctive words) does not fire
    assert _prose_describes("The realized-vs-implied volatility scatter",
                            ["implied volatility has been calm all year"]) is False


# ── Pass 1 — dashboard badge suppressed only on a clear contradiction with the body ────────────────
def test_badge_suppressed_only_on_clear_contradiction():
    from render.infographic.families.regime_dashboard import _badge_contradicts_body
    body = ("Inflation has fallen to 2.39%, near the bottom of its recent range; its momentum has bled "
            "out and decelerated hard. The labour market is still tight and vacancies remain high.")
    assert _badge_contradicts_body("CPI", "elevated", "rising", body)         # CPI↔inflation synonym, falls
    assert not _badge_contradicts_body("Vacancy rate", "elevated", "rising", body)  # body agrees → keep
    assert not _badge_contradicts_body("CPI", "elevated", "rising", "cpi feeds the model")  # ambiguous
    # proximity: a concept mention with the opposite cue FAR away (different clause) must NOT fire
    far = "Inflation is the model's key input. Separately, oil prices have fallen sharply this month."
    assert not _badge_contradicts_body("CPI", "elevated", "rising", far)


# ── D — one concept → one number ──────────────────────────────────────────────────────────────────
def test_concept_registry_collapses_same_concept_keeps_distinct_apart():
    from render.infographic.from_persona import _concept_registry
    numbers = {
        "variance_risk_premium.implied_vol_pct": _N("variance_risk_premium.implied_vol_pct", 15.57, "%"),
        "garch_volatility.implied_vol_pct": _N("garch_volatility.implied_vol_pct", 16.73, "%"),
        "funding_quality.ig_all_in_pct": _N("funding_quality.ig_all_in_pct", 5.86, "%"),
        "funding_quality.hy_all_in_pct": _N("funding_quality.hy_all_in_pct", 8.10, "%"),
    }
    meanings = {k: "" for k in numbers}
    runs = {"variance_risk_premium": {"meta": {"name": "VRP"}, "history": [0] * 252},
            "garch_volatility": {"meta": {"name": "GARCH"}, "history": [0] * 100},
            "funding_quality": {"meta": {"name": "Funding"}, "history": [0] * 500}}
    _concepts, canon = _concept_registry(numbers, meanings, template_keys=set(), runs=runs)
    iv = [k for k in numbers if "implied_vol" in k]
    assert canon[iv[0]] == canon[iv[1]]                               # two implied-vols → one
    assert canon[iv[0]] == "variance_risk_premium.implied_vol_pct"    # the longer-history model wins
    assert canon["funding_quality.ig_all_in_pct"] != canon["funding_quality.hy_all_in_pct"]  # IG≠HY


# ── A — semantic placement, robust to redraft drift ───────────────────────────────────────────────
def _cb_fixture():
    from render.writer import Outline, SectionPlan, SectionDraft
    outline = Outline(headline="H", standfirst="s", pivot="p", open_close="o", sections=[
        SectionPlan(heading="The real policy rate", thesis="t", model_id="real_rate",
                    chart_ids=["Real rate vs r*"]),
        SectionPlan(heading="Labour-market tightness", thesis="t", model_id="beveridge_curve",
                    chart_ids=["V/U tightness"]),
        SectionPlan(heading="The term spread", thesis="t", model_id="term_spread",
                    chart_ids=["Term spread 10y-3m"]),
    ])
    # writer REORDERS + RENAMES the sections and adds Phillips-curve prose
    filled = [
        SectionDraft(heading="Labour-market tightness and the Beveridge curve",
                     prose="the beveridge curve jumped outward; the phillips curve is the punchline."),
        SectionDraft(heading="By the real-rate yardstick", prose="the real_rate stance drained to neutral"),
        SectionDraft(heading="What the market prices — the term spread", prose="term_spread positive again"),
    ]
    ci = {
        "Real rate vs r*": {"model_id": "real_rate", "insight": ""},
        "V/U tightness": {"model_id": "beveridge_curve", "insight": ""},
        "Term spread 10y-3m": {"model_id": "term_spread", "insight": ""},
        "The Phillips curve (unemployment vs inflation)": {"model_id": "phillips_curve", "insight": ""},
    }
    return outline, filled, ci


def test_charts_bind_to_the_section_that_reads_them_after_reorder():
    from render.writer import _bind_sections, _place_charts
    outline, filled, ci = _cb_fixture()
    bindings = _bind_sections(outline, filled, ci)
    _place_charts(bindings, ci)
    for sec, b in zip(filled, bindings):
        sec.chart_ids = list(b.chart_ids)
    labour = next(s for s in filled if "Labour" in s.heading)
    realr = next(s for s in filled if "yardstick" in s.heading)
    assert "V/U tightness" in labour.chart_ids          # not one section off, as the old i%len did
    assert "Real rate vs r*" in realr.chart_ids


# ── B — completeness, high precision ──────────────────────────────────────────────────────────────
def test_named_exhibit_is_built_but_topic_mention_is_not():
    from render.writer import _bind_sections, _reconcile_prose_charts, _place_charts, _prose_names
    outline, filled, ci = _cb_fixture()
    full_text = " ".join(s.prose for s in filled)
    bindings = _bind_sections(outline, filled, ci)
    _reconcile_prose_charts(full_text, bindings, {"chart_index": ci}, conn=None, mat={"id": "cb", "runs": {}})
    _place_charts(bindings, ci)
    for sec, b in zip(filled, bindings):
        sec.chart_ids = list(b.chart_ids)
    built = {c for s in filled for c in s.chart_ids}
    assert "The Phillips curve (unemployment vs inflation)" in built    # promised → built
    # a chart whose TOPIC is discussed but which is not named as an exhibit must NOT be pulled
    assert not _prose_names("Reaction function — Taylor vs actual", {"model_id": "reaction_function"},
                            "the reaction function shows the fed is behind and the output gap is wide")


# ── C — dashboard from the finished article ───────────────────────────────────────────────────────
def test_dashboard_reads_the_article_and_falls_back_to_template():
    from render.infographic.from_persona import dashboard_thesis, dashboard_read, dashboard_tile_keys
    numbers = {"m.cpi_pct": _N("m.cpi_pct", 2.39, "%"), "m.taylor_pct": _N("m.taylor_pct", 5.31, "%"),
               "m.stance_pp": _N("m.stance_pp", 0.18, "pp"), "m.prob": _N("m.prob", 0.16, "")}
    mat = {"p": {"summary_template": "Template thesis. Stale template read."}, "numbers": numbers,
           "salient": ["m.cpi_pct", "m.taylor_pct", "m.stance_pp", "m.prob"], "canon": {}}
    article = {"exec_summary": "The Fed is behind its own rule. Inflation has fallen. "
                               "The gap is the bet, and it cannot be hedged both ways.",
               "cited_keys": ["m.stance_pp", "m.cpi_pct"]}
    assert "rule" in dashboard_thesis(mat, article).lower()
    assert "hedged" in dashboard_read(mat, article).lower()
    tk = dashboard_tile_keys(mat, article, n=4)
    assert tk[:2] == ["m.stance_pp", "m.cpi_pct"]            # the numbers the PROSE cited lead the tiles
    assert "template" in dashboard_thesis(mat, None).lower()  # clean fallback when no article
